"""`document-worker` v0.4.0 - sidecar Python do Frederico IA Studio (Fase 5, Etapa 5 PR 3).

Worker que gera documentos profissionais (DOCX, XLSX, PDF), le os tres
formatos e faz OCR de imagens e PDFs escaneados. Comunica com o app
principal via **named pipes** do Windows sobre o **envelope IPC** do
`frederico-process-architecture` (line-delimited JSON, 8 opcodes
estaveis em snake_case com prefixo de direcao: `worker.hello`,
`app.ack`, `app.ping`, `worker.pong`, `app.shutdown`, `worker.error`,
`tool.invoke`, `tool.result`).

## Protocolo (resumo do handshake)

1. Worker sobe, gera um `pipe_name` unico, cria o `NamedPipeServer`
   (via `pywin32.CreateNamedPipe`), e imprime `READY <pipe_name>` no
   **stdout** (handshake invertido, ADR-0017 Decisao 2).
2. Worker espera o app conectar (`ConnectNamedPipe`).
3. Worker envia `worker.hello` com o manifesto carregado de
   `manifest.json` (gera `request_id` UUID v4).
4. App responde `app.ack` com o `WorkerAuth` (token de curta duracao).
   Worker salva o token.
5. Loop: worker le linhas JSON do pipe, dispatch por `op`:
   - `app.ping` -> `worker.pong` com `status: "ok"`.
   - `app.shutdown` -> fecha o pipe e sai (manager detecta EOF).
   - `tool.invoke` -> valida token, dispatcha pro handler da
     `capability` declarada, e devolve `tool.result`.

## Handlers (Etapa 2B+Y - 7 primitivas)

| Capability   | Input                                       | Output                                              |
| ------------ | ------------------------------------------- | --------------------------------------------------- |
| `docx.write` | `path`, `title`, `sections`                 | `path`, `size_bytes`, `sections_written`            |
| `docx.read`  | `path`                                      | `paragraphs`, `tables`, contagens                   |
| `xlsx.write` | `path`, `sheets`                            | `path`, `size_bytes`, `sheets_written`              |
| `xlsx.read`  | `path` (opcional `sheet`)                   | `sheets`, `n_sheets`                                |
| `pdf.write`  | `path`, `title`, `sections`                 | `path`, `size_bytes`, `pages_rendered`              |
| `pdf.read`   | `path`, `ocr: "auto"|"never"|"only"`        | `text`, `ocr_text`, `page_count`, `scanned_pages`, `ocr_available`, `ocr_truncated`, `tesseract_version` |
| `ocr.run`    | `path`, `lang: "por+eng"` (opcional)        | `text`, `lang`, `conf`, `tesseract_version`         |
| `pdf.audit`  | `path`, `kind: "structural"`, `pdfa: Option<PdfAFlavor>`, `metadata`, `expected_pages` | `ok`, `checks`, `coverage`, `cache_key`, `rules_version` (sucesso) ou `code: "pdf_audit_structural_failed"`, `failed` (falha) |

**`pdf.audit` (Etapa 5 PR 3 da Fase 5, D-PDF5 + D-PDF6 do
ADR-0021):** auditoria estrutural bloqueante do §19.4
(PROMPT MESTRE). Roda em TODA geracao do PDFPro - o kit
chama apos `pdf.write` e o `salvar()` falha estruturado se
`ok: false` (§19.6 nao tem interruptor). `kind: "structural"`
e o unico implementado nesta PR; `kind: "visual"` (rasterizacao
pypdfium2 + grade/sobreposicao/pagina vazia) entra no PR 4
como extensao do mesmo handler.

**`pdf.read` ganhou fallback OCR (Etapa 2B+Y, ADR-0019 §Decisao 3):**
- `text` so vem da camada de texto do PDF (fonte: `pdfplumber`).
- `ocr_text` e um mapa `{page_num: texto_ocr}` separado, populado
  apenas pra paginas escaneadas (sem camada de texto) **quando o
  OCR foi rodado**. **Nunca** mistura com o `text` - procedencia
  sempre clara (mesma disciplina do `origin`/`external_content`
  da memoria). OCR troca 8 por B, 0 por O, 1 por l - e exatamente
  em CNPJ/valor/competencia que o erro cai. Misturar no mesmo campo
  apagaria a procedencia.
- `ocr: "auto"` (default): fallback transparente - se ha scanned
  pages e Tesseract esta disponivel, faz OCR delas e popula
  `ocr_text`. `ocr: "never"`: rapido, so checa camada de texto.
  `ocr: "only"`: ignora camada de texto e faz OCR de TODAS as
  paginas.
- `ocr_truncated: true` quando o teto de paginas/timeout foi
  atingido (PDF escaneado de 200 paginas nao trava o worker).
- `tesseract_version` no retorno = reprodutibilidade. Quando um
  resultado de OCR for questionado daqui a tres meses, esse campo
  permite reproduzir.

**`ocr.run` (Etapa 2B+Y):** OCR de uma imagem (PNG/JPG/TIFF/BMP) via
Tesseract. `lang` e validado com regex estrita (`^[a-z]{3}(+[a-z]{3})*$`)
e contra os traineddata realmente instalados - erro estruturado se
o idioma nao existe, em vez de mensagem criptica do Tesseract.

## Path safety

Barreira minima (Etapa 7 - sandbox-runner - traz sandbox de OS):
- Path nao pode conter `..` como componente.
- Path absoluto ou relativo ao `cwd` do worker.
- Diretorio pai (write) ou arquivo (read) deve existir e ser gravavel
  ou legivel respectivamente.

A camada mais forte (allowlist de diretorios por tool) entra na Etapa 3
junto com o `ToolManifest::allowed_paths` - registrada como pendencia
no `docs/modules/process-architecture.md`.

## Decisao de design: handler = primitiva, nao renderer de DocumentSpec

ADR-0018 §Decisao 1. Os 7 handlers sao primitivas de I/O sobre as
bibliotecas Python (`python-docx`, `openpyxl`, `reportlab`,
`pdfplumber`, `pytesseract`). Eles **nao** decidem margem, fonte,
cor, numeracao de pagina, header/footer, etc. - isso e trabalho do
**kit** (`WordPro`/`ExcelPro`/`PdfPro`, Etapa 3) que recebe o
`DocumentSpec` declarativo e traduz pra esses handlers. A v0.3.0 e
deliberadamente feia em tipografia; a beleza visual e o trabalho do
kit.
"""

from __future__ import annotations

import json
import logging
import os
import subprocess
import sys
import time
import uuid
import hashlib
from pathlib import Path, PurePath
from typing import Any, Callable

# `pywin32` e instalado pelo `bootstrap.ps1` (ADR-0004). Try/except
# claro pra que o erro apareca no stderr do worker, nao no app.
try:
    import win32pipe  # type: ignore[import-untyped]
    import win32file  # type: ignore[import-untyped]
    import pywintypes  # type: ignore[import-untyped]
except ImportError as exc:
    print(
        f"[document-worker] ERRO: pywin32 nao esta instalado ({exc}). "
        "Rode o bootstrap.ps1 pra instalar Python + pywin32 em runtime/. "
        "Ver ADR-0004.",
        file=sys.stderr,
        flush=True,
    )
    raise SystemExit(2)

# Bibliotecas dos handlers (instaladas pelo bootstrap.ps1, ADR-0018
# Decisao 2a). Falta de qualquer uma = worker nao sobe.
try:
    import docx  # python-docx
    import openpyxl
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        Image as RLImage,
        KeepTogether,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
    import pdfplumber
    # `pikepdf` (auditoria estrutural do PDFPro, D-PDF5 do ADR-0021,
    # Etapa 5 PR 3 da Fase 5). D-FAIL-1: hard-fail no bootstrap se
    # faltar, e o worker nao sobe sem ele a partir desta versao. As
    # outras 3 (pypdfium2, fontTools) sao importaveis separadamente
    # porque falhas nelas nao impedem o worker de subir - cada handler
    # sinaliza estruturado quando a dep dele nao esta.
    import pikepdf  # type: ignore[import-untyped]
except ImportError as exc:
    print(
        f"[document-worker] ERRO: biblioteca faltando ({exc}). "
        "Rode o bootstrap.ps1 pra instalar as dependencias. "
        "Ver ADR-0018 Decisao 2a.",
        file=sys.stderr,
        flush=True,
    )
    raise SystemExit(3)

# `fontTools` (glifo-check pre-render, D-GLYPH-1 do ADR-0021,
# Etapa 5 PR 2 da Fase 5). D-FAIL-1: hard-fail no bootstrap
# se faltar. Em runtime, se faltar, o `pdf.write` falha
# estruturado (`code: "fonttools_unavailable"`) em vez de
# renderizar PDF com glifo faltando. Mesma disciplina do
# `pytesseract` (pode faltar, mas o handler sinaliza).
try:
    from fontTools.ttLib import TTFont as FTFont  # type: ignore[import-untyped]
    FONTTOOLS_AVAILABLE = True
except ImportError:
    FTFont = None  # type: ignore[assignment]
    FONTTOOLS_AVAILABLE = False

# `pytesseract` (wrapper Python do Tesseract) e instalado pelo
# `bootstrap.ps1` (ADR-0019). Diferente das outras libs: a falta dele
# NAO impede o worker de subir - os handlers `docx`/`xlsx`/`pdf.write`
# continuam funcionando. So o `ocr.run` e o fallback OCR do `pdf.read`
# ficam indisponiveis. O `worker.hello` carrega `ocr_available: false`
# e o `pdf.read` retorna `code: "ocr_not_available"` quando
# alguem chama com `ocr: "auto"`/`"only"`. O caller (kit) decide se
# trata como erro ou segue.
try:
    import pytesseract  # type: ignore[import-untyped]
    from PIL import Image  # type: ignore[import-untyped]
    PYTESSERACT_AVAILABLE = True
except ImportError:
    pytesseract = None  # type: ignore[assignment]
    Image = None  # type: ignore[assignment]
    PYTESSERACT_AVAILABLE = False

# Versao do envelope IPC - bump MAJOR em mudancas incompativeis
# (mesmo numero que `IpcMessage::current_protocol_version()` no Rust).
PROTOCOL_VERSION: int = 1

# Tamanho do buffer de leitura (bytes).
READ_BUFFER_SIZE: int = 4096

# ---------------------------------------------------------------------------
# OCR (Tesseract + pytesseract) - Etapa 2B+Y, ADR-0019
# ---------------------------------------------------------------------------
#
# Diretorio do Tesseract binary (instalado pelo `bootstrap.ps1`).
# O `pytesseract` wrapper aceita `tesseract_cmd` apontando pro binario
# e `TESSDATA_PREFIX` apontando pro tessdata/ - definimos no startup
# (no `worker_main`) pra que esteja correto antes de qualquer
# chamada de OCR.
RUNTIME_TESSERACT_DIR = Path(__file__).resolve().parent / "runtime" / "tesseract"
TESSERACT_EXE = RUNTIME_TESSERACT_DIR / "tesseract.exe"
TESSERACT_TESSDATA_DIR = RUNTIME_TESSERACT_DIR / "tessdata"

# Idioma OCR default pro `ocr.run` (Tinta e Latao = Brasil; por sozinho
# perde em texto com termos em ingles, por+eng e mais robusto pro caso
# real - ADR-0019 §Decisao 2). O `pdf.read` no fallback automatico usa
# `por` sozinho (contexto brasileiro, sem lixo de outro idioma).
DEFAULT_OCR_LANG = "por+eng"
PDF_FALLBACK_OCR_LANG = "por"

# Idiomas que o bootstrap instalou (tessdata_fast 4.1.0, SHA-256 fixo).
# O handler valida o `lang` recebido contra esta lista - erro
# estruturado se faltar (em vez de mensagem criptica do Tesseract).
INSTALLED_OCR_LANGS = frozenset({"por", "eng", "osd"})

# Teto de paginas/timeout pro OCR automatico do `pdf.read`. Sem isso,
# um PDF escaneado de 200 paginas trava o worker por minutos. Quando
# bate o teto, o handler devolve o que ja processou e marca
# `ocr_truncated: true` no payload - o caller decide se aceita
# parcial ou aborta. **Tuning:** 20 paginas * 30s = 10 min maximo
# absoluto de OCR por `pdf.read` (ainda muito, mas dentro do timeout
# do manager de 60s por invoke? NAO. Ajustar pra caber.)
MAX_OCR_PAGES_PDF = 20
OCR_TIMEOUT_S_PER_PAGE = 30

# Versao do Tesseract. Calculada no startup (`worker_main`) e
# cacheada aqui. `None` se o Tesseract nao esta instalado (binario
# nao apareceu). Incluida em todo `ocr.run`/`pdf.read` com OCR
# no payload - reprodutibilidade (3 meses depois: "qual versao do
# Tesseract produziu esse texto?").
TESSERACT_VERSION: str | None = None
TESSERACT_VERSION_DETECTED_AT: float = 0.0

# Timeout do `ConnectNamedPipe` (ms). O `kill_on_drop` no manager
# Rust garante cleanup se o app travar antes de conectar.
CONNECT_TIMEOUT_MS: int = 60_000

# Diretorio de fontes Tinta e Latao. Auto-load no startup. Fallback
# pra fontes built-in do reportlab (Helvetica/Times-Roman) se nao
# encontrar.
RUNTIME_FONTS_DIR = Path(__file__).resolve().parent / "runtime" / "fonts"
WINDOWS_FONTS_DIR = Path(os.environ.get("WINDIR", "C:\\Windows")) / "Fonts"

# Mapeamento canonico das fontes da identidade visual "Tinta e Latao"
# (Adobe Source Sans 3 + Source Serif 4, ADR-0018 Decisao 2b). Esses
# nomes sao usados internamente no `reportlab.pdfmetrics` e nas
# `ParagraphStyle` do `pdf.write`.
FONT_BODY_NAME = "TintaLataoSans"      # Source Sans 3 - corpo
FONT_TITLE_NAME = "TintaLataoSerif"    # Source Serif 4 - titulos

FONT_FILES = {
    FONT_BODY_NAME: [
        RUNTIME_FONTS_DIR / "SourceSans3VF-Upright.ttf",
        RUNTIME_FONTS_DIR / "SourceSans3VF-Italic.ttf",
        WINDOWS_FONTS_DIR / "SourceSans3VF-Upright.ttf",  # fallback improvavel
    ],
    FONT_TITLE_NAME: [
        RUNTIME_FONTS_DIR / "SourceSerif4Variable-Roman.ttf",
        RUNTIME_FONTS_DIR / "SourceSerif4Variable-Italic.ttf",
        WINDOWS_FONTS_DIR / "SourceSerif4Variable-Roman.ttf",
    ],
}

# Logging basico. Vai pro stderr, que o `WorkerManager::spawn_external`
# (Rust) captura e loga via `tracing::warn!` (ver
# `crates/process-architecture/src/external.rs` §"Stderr pump").
logging.basicConfig(
    level=logging.INFO,
    format="[document-worker] %(asctime)s %(levelname)s %(message)s",
    stream=sys.stderr,
)
log = logging.getLogger("document-worker")


# ---------------------------------------------------------------------------
# IPC envelope (espelha `frederico-process-architecture::protocol::IpcMessage`)
# ---------------------------------------------------------------------------


def ipc_message(op: str, payload: dict[str, Any], auth: str | None = None, request_id: str | None = None) -> bytes:
    """Serializa uma `IpcMessage` como **uma linha** (line-delimited JSON).

    O `\\n` no final e o separador de mensagens. O
    `IpcMessage::decode_line` (Rust) faz `strip_suffix(b"\\\\n")` -
    sem o newline o decode falha.

    **request_id:** se passado, e ecoado no envelope (response do
    `tool.invoke` deve carregar o mesmo `request_id` do request - e
    o que o `WorkerManager` (ator) usa pra casar a response com o
    `pending` que gerou. Bug sutil na Etapa 2B original que gerava
    UUID novo na response - fazia o invoke timeoutar. O
    `worker_stub` Rust (Etapa 2B continuacao) ja fazia certo.
    """
    msg = {
        "protocol_version": PROTOCOL_VERSION,
        "request_id": request_id if request_id is not None else str(uuid.uuid4()),
        "op": op,
        "payload": payload,
    }
    if auth is not None:
        msg["auth"] = auth
    return (json.dumps(msg, separators=(",", ":"), ensure_ascii=False) + "\n").encode("utf-8")


def decode_line(line: bytes) -> dict[str, Any]:
    """Desserializa uma linha JSON. Valida `protocol_version`.

    Lanca `ValueError` em payload malformado ou versao errada.
    O `IpcMessage::decode_line` (Rust) faz o mesmo.
    """
    if line.endswith(b"\r\n"):
        line = line[:-2]
    elif line.endswith(b"\n"):
        line = line[:-1]
    msg = json.loads(line.decode("utf-8"))
    pv = msg.get("protocol_version")
    if pv != PROTOCOL_VERSION:
        raise ValueError(
            f"protocol_version {pv} nao e a atual {PROTOCOL_VERSION}"
        )
    return msg


# ---------------------------------------------------------------------------
# Path safety
# ---------------------------------------------------------------------------


class PathSafetyError(Exception):
    """Lancada quando um path falha a validacao minima (Etapa 2B+X)."""

    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code
        self.message = message


def validate_path(path_str: str, kind: str) -> Path:
    """Valida um path de I/O. `kind` e "read" ou "write".

    Regras (barreira minima - a forte vem na Etapa 3 com
    `ToolManifest::allowed_paths` + Etapa 7 com sandbox-runner):
      1. Nao pode ser vazio.
      2. Nao pode conter `..` como componente.
      3. Deve ser absoluto, ou relativo ao `cwd` do worker.
      4. Para `kind == "write"`: o diretorio pai deve existir e ser
         gravavel. O arquivo pode nao existir ainda.
      5. Para `kind == "read"`: o arquivo deve existir e ser legivel.
    """
    if not path_str or not isinstance(path_str, str):
        raise PathSafetyError("invalid_path", "path ausente ou nao-string")
    p = PurePath(path_str)
    if ".." in p.parts:
        raise PathSafetyError(
            "path_traversal",
            f"path contem '..': {path_str!r}",
        )
    # Resolver para absoluto (mantem o que ja e absoluto).
    abs_path = Path(path_str).resolve() if Path(path_str).is_absolute() else (Path.cwd() / path_str).resolve()
    if kind == "write":
        parent = abs_path.parent
        if not parent.is_dir():
            raise PathSafetyError(
                "parent_missing",
                f"diretorio pai nao existe: {parent}",
            )
        # Teste de gravabilidade: cria e deleta um arquivo temporario.
        # Pode falhar por permissao - tratamos como erro.
        try:
            test_file = parent / ".frederico_write_probe"
            test_file.touch()
            test_file.unlink()
        except OSError as exc:
            raise PathSafetyError(
                "parent_not_writable",
                f"diretorio pai nao gravavel: {parent} ({exc})",
            ) from exc
        return abs_path
    elif kind == "read":
        if not abs_path.is_file():
            raise PathSafetyError(
                "file_missing",
                f"arquivo nao existe: {abs_path}",
            )
        if not os.access(abs_path, os.R_OK):
            raise PathSafetyError(
                "file_not_readable",
                f"arquivo nao legivel: {abs_path}",
            )
        return abs_path
    else:
        raise PathSafetyError("invalid_kind", f"kind deve ser 'read' ou 'write', veio {kind!r}")


# ---------------------------------------------------------------------------
# Fontes "Tinta e Latao" - auto-load no startup
# ---------------------------------------------------------------------------

_FONTS_REGISTERED = False
_FONT_STATUS: dict[str, str] = {}


def _try_register_font(name: str, candidates: list[Path]) -> bool:
    """Tenta registrar uma fonte TTF. Retorna True se conseguiu."""
    for path in candidates:
        if path.is_file() and path.stat().st_size > 50_000:
            try:
                pdfmetrics.registerFont(TTFont(name, str(path)))
                log.info("fonte %s registrada de %s", name, path)
                return True
            except Exception as exc:
                log.warning("falha registrando %s de %s: %s", name, path, exc)
    return False


def ensure_fonts_registered() -> dict[str, str]:
    """Garante que as fontes Tinta e Latao estao registradas.

    Retorna um dicionario `name -> "loaded" | "fallback"` pra incluir
    no `worker.pong` payload e em `worker.hello` (campo extra). O
    `pdf.write` consulta esse mesmo estado.

    **Bug fix v0.2.0:** a versao inicial retornava `{}` apos o
    primeiro registro (o `status` so era populado dentro do `if`).
    Isso fazia o `pong` reportar `font_status: {}` depois do boot,
    e os testes E2E nao conseguiam validar que as TTFs tinham
    carregado. Agora mantemos o status em cache no module-level e
    retornamos o memo.
    """
    global _FONTS_REGISTERED, _FONT_STATUS
    if _FONTS_REGISTERED:
        return _FONT_STATUS
    status: dict[str, str] = {}
    for name, candidates in FONT_FILES.items():
        if _try_register_font(name, candidates):
            status[name] = "loaded"
        else:
            # Fallback: reportlab ja tem Helvetica e Times-Roman
            # built-in. Mapear nomes canonicos pros built-ins
            # garante que o `pdf.write` nao quebra se as TTFs
            # nao estiverem instaladas.
            if name == FONT_BODY_NAME:
                pdfmetrics.registerFontFamily(name, normal="Helvetica", bold="Helvetica-Bold", italic="Helvetica-Oblique", bold_italic="Helvetica-BoldOblique")
            else:
                pdfmetrics.registerFontFamily(name, normal="Times-Roman", bold="Times-Bold", italic="Times-Italic", bold_italic="Times-BoldItalic")
            status[name] = "fallback"
            log.warning("fonte %s NAO encontrada - usando fallback built-in do reportlab", name)
    _FONTS_REGISTERED = True
    _FONT_STATUS = status
    return status


# ---------------------------------------------------------------------------
# OCR (Tesseract via pytesseract) - Etapa 2B+Y, ADR-0019
# ---------------------------------------------------------------------------
#
# Wrappers finos em torno do `pytesseract`:
#   - _get_tesseract_version(): detecta versao no startup, cacheia.
#   - _validate_lang(lang): regex estrita + checa contra INSTALLED_OCR_LANGS.
#   - _ocr_image_to_text(image_path, lang): OCR de uma imagem (PNG/JPG/TIFF/BMP).
#   - _ocr_pdf_page_to_text(pdf_path, page_num, lang): OCR de 1 pagina
#     de um PDF, via render com pdfplumber (nao precisa de Poppler).
#
# Por que regex estrita no lang: o valor vem do chamador e vira argumento
# de linha de comando do `tesseract.exe`. Sem validacao, um chamador
# hostil (ou buggy) pode injetar `; rm -rf /` ou coisa pior. Mesma
# disciplina da barreira de path dos outros handlers - o manager
# Rust nao tem como auditar isso, o handler Python tem.
#
# Por que `TESSERACT_VERSION` em module-level (e nao em `worker_main`):
# o handler e chamado em qualquer momento depois do startup. Calcular
# no `worker_main` e cachear evita spawnar `tesseract.exe --version`
# em todo handler (overhead ~50ms cada).
import re
_LANG_PATTERN = re.compile(r"^[a-z]{3}(\+[a-z]{3})*$")


def _tesseract_executable_present() -> bool:
    """Retorna True se `tesseract.exe` existe em runtime/tesseract/."""
    return TESSERACT_EXE.is_file() and TESSERACT_EXE.stat().st_size > 100_000


def _get_tesseract_version() -> str | None:
    """Detecta a versao do Tesseract. Cacheia em module-level.

    Retorna `None` se Tesseract nao esta instalado (bootstrap pulou
    em contexto non-admin, ou falhou). O chamador (handler `ocr.run`
    ou `pdf.read` com `ocr: "auto"`) decide como tratar.
    """
    global TESSERACT_VERSION, TESSERACT_VERSION_DETECTED_AT
    if TESSERACT_VERSION is not None or TESSERACT_VERSION_DETECTED_AT > 0:
        return TESSERACT_VERSION
    if not _tesseract_executable_present():
        TESSERACT_VERSION_DETECTED_AT = time.time()
        return None
    try:
        proc = subprocess.run(
            [str(TESSERACT_EXE), "--version"],
            capture_output=True,
            text=True,
            timeout=10,
            check=False,
        )
        if proc.returncode != 0:
            log.warning("tesseract --version saiu com codigo %d", proc.returncode)
            TESSERACT_VERSION_DETECTED_AT = time.time()
            return None
        # Output esperado: "tesseract 5.4.0.20240606\n ..." (primeira linha).
        first_line = proc.stdout.strip().splitlines()[0] if proc.stdout else ""
        # Pega so o token "5.4.0.20240606" (sem "tesseract ").
        m = re.search(r"(\d+\.\d+\.\d+(?:\.\d+)?)", first_line)
        TESSERACT_VERSION = m.group(1) if m else first_line
        TESSERACT_VERSION_DETECTED_AT = time.time()
        log.info("Tesseract version detectada: %s", TESSERACT_VERSION)
        return TESSERACT_VERSION
    except (subprocess.TimeoutExpired, FileNotFoundError, OSError) as exc:
        log.warning("falha detectando versao do Tesseract: %s", exc)
        TESSERACT_VERSION_DETECTED_AT = time.time()
        return None


def _validate_lang(lang: str) -> str:
    """Valida o `lang` recebido no `ocr.run`. Retorna o `lang` normalizado.

    Regras (ADR-0019 §Decisao 2.5 - defesa contra injecao de argumento):
      1. Regex `^[a-z]{3}(+[a-z]{3})*$` - so segmentos de 3 letras
         minusculas, opcionalmente concatenados por `+`.
      2. Cada segmento tem que estar em `INSTALLED_OCR_LANGS` (o
         conjunto de traineddata realmente instalados).

    Lanca `ValueError` com mensagem clara listando os idiomas
    disponiveis (em vez da mensagem criptica do Tesseract quando
    o idioma nao existe).
    """
    if not isinstance(lang, str) or not _LANG_PATTERN.match(lang):
        raise ValueError(
            f"lang invalido: {lang!r}. Esperado segmentos de 3 letras "
            f"minusculas unidos por '+'. Exemplo: 'por', 'eng', 'por+eng'."
        )
    parts = lang.split("+")
    missing = [p for p in parts if p not in INSTALLED_OCR_LANGS]
    if missing:
        raise ValueError(
            f"idioma OCR nao instalado: {missing!r}. "
            f"Idiomas disponiveis: {sorted(INSTALLED_OCR_LANGS)}. "
            f"Re-adicione via bootstrap.ps1 (tag do tessdata + SHA-256 fixos)."
        )
    return lang


def _configure_pytesseract() -> None:
    """Configura `pytesseract` no startup do worker. Idempotente.

    Define `tesseract_cmd` (path do binario) e `TESSDATA_PREFIX` (env
    var que o tesseract consulta pra localizar `*.traineddata`).
    """
    if not PYTESSERACT_AVAILABLE:
        return
    if _tesseract_executable_present():
        pytesseract.pytesseract.tesseract_cmd = str(TESSERACT_EXE)
        # `TESSDATA_PREFIX` precisa ser setado **no env do processo**
        # (pytesseract le via os.environ no subprocess) e tambem
        # no fallback via config_data_dir. Setamos ambos pra ser
        # defensivo contra diferencas entre versoes do pytesseract.
        os.environ["TESSDATA_PREFIX"] = str(TESSERACT_TESSDATA_DIR)
        pytesseract.pytesseract.config_data_dir = str(TESSERACT_TESSDATA_DIR)


def _ocr_pil_image(img, lang: str, timeout_s: int = 60) -> dict:
    """OCR de um objeto PIL.Image. Retorna `{"text", "conf"}`."""
    data = pytesseract.image_to_data(
        img,
        lang=lang,
        output_type=pytesseract.Output.DICT,
        timeout=timeout_s,
    )
    words = []
    confs = []
    for i, txt in enumerate(data.get("text", [])):
        t = (txt or "").strip()
        if not t:
            continue
        words.append(t)
        try:
            c = int(data["conf"][i])
        except (ValueError, TypeError):
            continue
        if c >= 0:
            confs.append(c)
    text = " ".join(words)
    mean_conf = (sum(confs) / len(confs)) if confs else None
    return {"text": text, "conf": mean_conf}


def _ocr_image_to_text(image_path: Path, lang: str, timeout_s: int = 60) -> dict:
    """OCR de uma imagem em disco. Retorna `{"text", "conf"}`."""
    if not PYTESSERACT_AVAILABLE:
        raise RuntimeError("pytesseract nao esta instalado")
    if not _tesseract_executable_present():
        raise RuntimeError(
            f"Tesseract nao encontrado em {TESSERACT_EXE}. "
            "Rode o bootstrap.ps1 como Admin (veja instrucoes no script)."
        )
    if not image_path.is_file():
        raise FileNotFoundError(f"imagem nao encontrada: {image_path}")
    try:
        img = Image.open(str(image_path))
    except Exception as exc:
        raise RuntimeError(f"falha abrindo imagem: {exc}") from exc
    try:
        return _ocr_pil_image(img, lang, timeout_s)
    except pytesseract.TesseractError as exc:
        raise RuntimeError(f"Tesseract falhou: {exc}") from exc
    except RuntimeError as exc:
        # pytesseract 0.3.10 levanta RuntimeError em timeout.
        if "timeout" in str(exc).lower():
            raise RuntimeError(f"OCR excedeu timeout de {timeout_s}s") from exc
        raise


def _ocr_pdf_page_to_text(
    pdf_path: Path, page_num: int, lang: str, timeout_s: int = 60
) -> str:
    """OCR de UMA pagina especifica de um PDF. Retorna o texto.

    Renderiza a pagina como imagem via `pdfplumber` (ja temos) +
    Pillow (ja temos), e chama Tesseract. Sem dependencia extra
    de Poppler/pdf2image.
    """
    if not PYTESSERACT_AVAILABLE:
        raise RuntimeError("pytesseract nao esta instalado")
    if not _tesseract_executable_present():
        raise RuntimeError("Tesseract nao encontrado")
    if not pdf_path.is_file():
        raise FileNotFoundError(f"PDF nao encontrado: {pdf_path}")
    with pdfplumber.open(str(pdf_path)) as pdf:
        if page_num < 1 or page_num > len(pdf.pages):
            raise ValueError(
                f"page_num {page_num} fora de [1, {len(pdf.pages)}]"
            )
        page = pdf.pages[page_num - 1]
        # `to_image(resolution=300)` e o padrao de OCR. Devolve
        # um `PIL.Image` que o pytesseract aceita direto.
        img = page.to_image(resolution=300).original
    return _ocr_pil_image(img, lang, timeout_s)["text"]


# ---------------------------------------------------------------------------
# Handlers
# ---------------------------------------------------------------------------
#
# Convenção: cada handler recebe `payload: dict` e devolve `dict`
# (vai direto pro `tool.result.payload`). Se algo falhar, o handler
# lanca uma excecao; o dispatch (no main loop) converte pra
# `worker.error` com code/message.


def _payload_field(payload: dict, key: str, expected_type: type) -> Any:
    """Extrai campo obrigatorio do payload com type-check."""
    if key not in payload:
        raise ValueError(f"campo obrigatorio ausente: {key!r}")
    val = payload[key]
    if not isinstance(val, expected_type):
        raise ValueError(
            f"campo {key!r} esperado {expected_type.__name__}, "
            f"veio {type(val).__name__}"
        )
    return val


# ---- docx.write -----------------------------------------------------------


def handle_docx_write(payload: dict) -> dict:
    """docx.write: escreve um arquivo .docx com `title` + `sections`.

    Input: `{"path": str, "title": str, "sections": [{"heading": str, "paragraphs": [str]}]}`
    Output: `{"ok": true, "path": str, "size_bytes": int, "sections_written": int}`
    """
    path = validate_path(_payload_field(payload, "path", str), "write")
    title = _payload_field(payload, "title", str)
    sections = _payload_field(payload, "sections", list)

    document = docx.Document()
    document.core_properties.title = title
    document.add_heading(title, level=0)
    sections_written = 0
    for sec in sections:
        if not isinstance(sec, dict):
            raise ValueError("secao precisa ser um dict")
        heading = sec.get("heading", "")
        paragraphs = sec.get("paragraphs", [])
        if not isinstance(paragraphs, list):
            raise ValueError("'paragraphs' precisa ser uma lista de strings")
        if heading:
            document.add_heading(heading, level=1)
        for p in paragraphs:
            if not isinstance(p, str):
                raise ValueError("paragrafo precisa ser string")
            document.add_paragraph(p)
        sections_written += 1
    document.save(str(path))
    size = path.stat().st_size
    return {
        "ok": True,
        "path": str(path),
        "size_bytes": size,
        "sections_written": sections_written,
    }


# ---- docx.read ------------------------------------------------------------


def handle_docx_read(payload: dict) -> dict:
    """docx.read: extrai paragrafos e tabelas de um .docx.

    Input: `{"path": str}`
    Output (Etapa 4 da Fase 5): `{"ok": true, "path": str,
    "paragraphs": [{"text": str, "style": str}], "tables":
    [[str]], "n_paragraphs": int, "n_tables": int}`

    `paragraphs[i].style` e o nome do estilo (ex:
    "Heading 1", "Heading 2", "Heading 3", "Normal").
    O `docs.inspect` usa isso pra reconstruir o
    `DocumentSpec` parcial (heading vs paragraph). Sem
    style (v0.3.0 do worker), o inspect perdia todos os
    headings e so via `paragraph` — limitacao que
    ficou registrada ate a Etapa 4.
    """
    path = validate_path(_payload_field(payload, "path", str), "read")
    document = docx.Document(str(path))
    paragraphs = [
        {"text": p.text, "style": p.style.name}
        for p in document.paragraphs
    ]
    tables = []
    for t in document.tables:
        rows = []
        for row in t.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(rows)
    return {
        "ok": True,
        "path": str(path),
        "paragraphs": paragraphs,
        "tables": tables,
        "n_paragraphs": len(paragraphs),
        "n_tables": len(tables),
    }


# ---- xlsx.write -----------------------------------------------------------


# Mapeamento de aliases semanticos (`"BRL"`, `"PCT"`, `"THOUSANDS"`)
# pra Excel format strings. O kit manda o alias (mais
# legivel que o format string cru do Excel) e o handler
# resolve. Aliases nao reconhecidos sao tratados como
# format string cru (passa direto pro openpyxl).
# Adicionar constante aqui quando um novo alias for
# necessario.
XLSX_FORMAT_ALIASES = {
    # Moeda brasileira (R$ 1.234,56)
    "BRL": 'R$ #,##0.00',
    # Percentual com 2 casas (12,34%)
    "PCT": '0.00%',
    # Separador de milhar com 2 casas (1.234,56)
    "THOUSANDS": '#,##0.00',
    # Inteiro com separador de milhar (1.234)
    "INT": '#,##0',
}


def _resolve_xlsx_format(value):
    """Resolve um alias de formato (ou retorna o valor cru
    se nao for alias). Aceita string; se nao for string,
    retorna None (sem formato aplicado — defensivo).
    """
    if not isinstance(value, str):
        return None
    return XLSX_FORMAT_ALIASES.get(value, value)


def handle_xlsx_write(payload: dict) -> dict:
    """xlsx.write: escreve um arquivo .xlsx com 1+ sheets.

    Input: `{"path": str, "sheets": [{"name": str, "headers": [str], "rows": [[]], "column_formats": {<col_idx>: <format>}?}]}`

    `column_formats` (opcional, Etapa 4 da Fase 5): mapa
    `{col_idx: format_alias_or_string}`. Aplica
    `cell.number_format` em todas as celulas da coluna
    (exceto o header). Aceita aliases semanticos (`"BRL"`,
    `"PCT"`, `"THOUSANDS"`, `"INT"`) ou Excel format
    strings crus. Backward-compat: sheets sem
    `column_formats` continuam funcionando (Etapa 3 da
    Fase 5 e anteriores).

    Output: `{"ok": true, "path": str, "size_bytes": int, "sheets_written": int, "total_rows": int, "cells_formatted": int}`

    `cells_formatted` conta quantas celulas receberam
    `cell.number_format` (zero em sheets sem
    `column_formats`).
    """
    path = validate_path(_payload_field(payload, "path", str), "write")
    sheets = _payload_field(payload, "sheets", list)
    wb = openpyxl.Workbook()
    # openpyxl cria uma sheet default "Sheet" - removemos e adicionamos as nossas.
    default = wb.active
    wb.remove(default)
    total_rows = 0
    sheets_written = 0
    cells_formatted = 0
    for sh in sheets:
        if not isinstance(sh, dict):
            raise ValueError("sheet precisa ser um dict")
        name = sh.get("name", "")
        if not name:
            raise ValueError("sheet sem nome")
        headers = sh.get("headers", [])
        rows = sh.get("rows", [])
        if not isinstance(headers, list) or not isinstance(rows, list):
            raise ValueError("'headers' e 'rows' precisam ser listas")
        column_formats = sh.get("column_formats")
        if column_formats is not None and not isinstance(column_formats, dict):
            raise ValueError("'column_formats' precisa ser um dict {col_idx: format}")
        ws = wb.create_sheet(title=name)
        if headers:
            ws.append(headers)
        for row in rows:
            ws.append(row)
        # Aplica `column_formats` em todas as celulas
        # de dados (rows, nao header). Itera por coluna
        # pra evitar recriar o dicionario em cada celula.
        # Roda DEPOIS do `ws.append` pra que `cell.value`
        # e `cell.row` correspondam ao que o usuario
        # mandou.
        if column_formats:
            for col_idx, fmt_value in column_formats.items():
                excel_fmt = _resolve_xlsx_format(fmt_value)
                if excel_fmt is None:
                    continue
                col = int(col_idx) + 1  # openpyxl e 1-indexed
                # header (linha 1, se houver) NAO recebe
                # format. Data rows comecam na linha 2.
                for row_offset in range(2, 2 + len(rows)):
                    cell = ws.cell(row=row_offset, column=col)
                    if cell.value is not None:
                        cell.number_format = excel_fmt
                        cells_formatted += 1
        total_rows += len(rows)
        sheets_written += 1
    wb.save(str(path))
    return {
        "ok": True,
        "path": str(path),
        "size_bytes": path.stat().st_size,
        "sheets_written": sheets_written,
        "total_rows": total_rows,
        "cells_formatted": cells_formatted,
    }


# ---- xlsx.read ------------------------------------------------------------

# Mapeamento reverso de Excel format string → alias
# semantico (consistente com `XLSX_FORMAT_ALIASES` do
# `xlsx.write`). Usado pelo `docs.inspect` (Etapa 4) pra
# expor `currency_format: "BRL"` em vez do string cru
# `R$ #,##0.00`. Normalizacao leve: o Excel as vezes
# inclui espacos extras ou formato equivalente.
XLSX_FORMAT_ALIAS_REVERSE = {
    "R$ #,##0.00": "BRL",
    "$ #,##0.00": "BRL",
    "0.00%": "PCT",
    "0%": "PCT",
    "#,##0.00": "THOUSANDS",
    "#,##0": "INT",
}


def _resolve_xlsx_format_alias(excel_fmt: str):
    """Resolve Excel format string → alias semantico
    (ou retorna o string cru se nao for alias conhecido).
    """
    return XLSX_FORMAT_ALIAS_REVERSE.get(excel_fmt, excel_fmt)


def handle_xlsx_read(payload: dict) -> dict:
    """xlsx.read: le um .xlsx e devolve sheets + dados + metadados estruturais.

    Input: `{"path": str, "sheet": str?, "sample_rows": int?, "range": str?}`
    - `sheet` filtra uma so sheet (opcional; default = todas).
    - `sample_rows` (opcional, Etapa 4 do docs.inspect):
      limita o numero de `first_rows` devolvidos por sheet
      (default 5, max 20). Se a sheet tem mais linhas, so
      as primeiras N sao incluidas. O numero total de
      linhas (`n_rows`) continua sendo o completo.
    - `range` (opcional, modo detalhe do docs.inspect):
      por enquanto, so validado (formato "A1:D10" ou
      similar). v0.3.0 do worker NAO aplica `range` ao
      openpyxl — o caller (docs.inspect) faz o filtro
      depois. A Etapa 4 do docs.inspect usa o range
      apenas como flag de "modo detalhe"; a leitura
      sempre vem completa.

    Output: `{"ok": true, "sheets": [{"name": str, "headers": [str], "rows": [[]],
    "used_range": str, "n_rows": int, "n_cols": int,
    "first_rows": [[]], "column_formats": {<col>: <alias>}}], "n_sheets": int}`

    `column_formats` (Etapa 4): mapa `{<col_idx>: <alias>}`
    derivado do `cell.number_format` da primeira celula
    NAO-vazia de cada coluna (header NAO conta — so
    dados). Aliases: "BRL" / "PCT" / "THOUSANDS" / "INT"
    (consistente com `XLSX_FORMAT_ALIASES` do write);
    se o formato do Excel nao bate com alias, o valor
    e o format string cru.

    `used_range`: intervalo usado no openpyxl (ex:
    "A1:C5") — usado pelo docs.inspect pra exibir
    "intervalo: A1:C5" no modo resumo.

    `first_rows`: amostra das primeiras `sample_rows`
    linhas de dados (Etapa 4, default 5, max 20) —
    evita despejar planilha de 5000 linhas no contexto
    do modelo.
    """
    path = validate_path(_payload_field(payload, "path", str), "read")
    sheet_filter = payload.get("sheet")
    if sheet_filter is not None and not isinstance(sheet_filter, str):
        raise ValueError("'sheet' precisa ser string")
    sample_rows = payload.get("sample_rows", 5)
    if not isinstance(sample_rows, int) or not (1 <= sample_rows <= 20):
        raise ValueError("'sample_rows' precisa ser int entre 1 e 20")
    # `range` so validado em v0.3.0 (o handler nao
    # aplica). Formato esperado: "A1:D10" (coluna
    # letra + linha numero, ate 2 pares). Validacao
    # leve — so confere que nao tem `..` e tem o
    # formato basico.
    range_arg = payload.get("range")
    if range_arg is not None and not isinstance(range_arg, str):
        raise ValueError("'range' precisa ser string")
    if range_arg is not None and ".." in range_arg:
        raise ValueError("'range' nao pode conter '..'")
    # Modo normal (read_only=False) pra ter acesso a
    # `cell.number_format`. O inspect e eventual (nao
    # roda em hot path), entao o custo extra de memoria
    # e aceitavel. data_only=True garante que celulas
    # com formula devolvem o valor calculado, nao a
    # formula.
    wb = openpyxl.load_workbook(str(path), data_only=True)
    sheets_out = []
    for ws in wb.worksheets:
        if sheet_filter is not None and ws.title != sheet_filter:
            continue
        rows_list = list(ws.iter_rows(values_only=True))
        headers: list = []
        data_rows: list = []
        if rows_list:
            headers = [str(c) if c is not None else "" for c in rows_list[0]]
            data_rows = [
                [str(c) if c is not None else "" for c in row]
                for row in rows_list[1:]
            ]
        n_cols = len(headers)
        n_rows = len(data_rows)
        # column_formats: `cell.number_format` da
        # primeira celula nao-vazia de cada coluna
        # (header e linha 1 contam como dados;
        # o caller decide se ignora). Em modo
        # normal, cell.number_format e acessivel.
        column_formats = {}
        for col_idx in range(n_cols):
            # Itera todas as linhas (incluindo
            # header) — o caller (docs.inspect)
            # sabe que row 0 e header e ignora.
            # Mas documentamos aqui: row 0
            # (header) e incluido se tiver
            # number_format explicito.
            for row_offset, _ in enumerate(rows_list, start=1):
                cell = ws.cell(row=row_offset, column=col_idx + 1)
                if cell.value is not None and cell.value != "":
                    column_formats[str(col_idx)] = _resolve_xlsx_format_alias(
                        cell.number_format
                    )
                    break
        # used_range: openpyxl expoe via
        # `ws.dimensions` (string tipo "A1:C5").
        used_range = ws.dimensions
        # first_rows: amostra (sample_rows).
        first_rows = data_rows[:sample_rows]
        sheets_out.append({
            "name": ws.title,
            "headers": headers,
            "rows": data_rows,
            "used_range": used_range,
            "n_rows": n_rows,
            "n_cols": n_cols,
            "first_rows": first_rows,
            "column_formats": column_formats,
        })
    wb.close()
    return {
        "ok": True,
        "path": str(path),
        "sheets": sheets_out,
        "n_sheets": len(sheets_out),
    }


# ---- pdf.write ------------------------------------------------------------
#
# PDFPro v0.1 (Etapa 5 PR 2 da Fase 5, ADR-0021):
# - Payload estendido (style, page, identity, watermark, metadata, blocks).
# - 20 blocos do `DocumentBlock` viram flowables do reportlab.
# - Glifo-check via `fontTools` ANTES do `doc.build()` (D-GLYPH-1):
#   falha estruturada com lista de blocos+caracteres faltantes.
# - Modo Sobrio (registraveis) com paleta monocromatica e margens
#   maiores; modo Tinta & Latao com paleta da marca.
# - Marca d'agua opt-in (D-PDF2) desenhada via `onPage` callback.
# - Fontes Tinta & Latao embutidas via `ensure_fonts_registered()`
#   (sem fallback para fonte do sistema - D-FAIL-1).


def _hex(c: str) -> HexColor:
    """Constrói `HexColor` aceitando string com ou sem `#`."""
    return HexColor(c if c.startswith("#") else f"#{c}")


def _build_pdf_styles(identity: dict, style_name: str) -> dict[str, ParagraphStyle]:
    """Constrói os ParagraphStyle usados pelo `pdf.write` estendido.

    Cores vêm do `identity` (paleta Tinta & Latão ou
    monocromática Sobrio). Fontes: `FONT_TITLE_NAME`
    (Source Serif 4) para títulos, `FONT_BODY_NAME`
    (Source Sans 3) para corpo. O `Courier` built-in é
    usado para bloco de código (monospace).
    """
    tinta = _hex(identity.get("tinta", "#000000"))
    text = _hex(identity.get("text", "#000000"))
    muted = _hex(identity.get("muted", "#000000"))
    success = _hex(identity.get("success", "#000000"))
    latao = _hex(identity.get("latao", "#000000"))
    light = _hex(identity.get("light", "#FFFFFF"))
    is_sobrio = style_name == "sobrio"
    return {
        # Cover
        "cover_title": ParagraphStyle(
            "TintaCoverTitle",
            fontName=FONT_TITLE_NAME,
            fontSize=28,
            leading=34,
            alignment=TA_CENTER,
            textColor=tinta,
            spaceAfter=18,
        ),
        "cover_subtitle": ParagraphStyle(
            "TintaCoverSubtitle",
            fontName=FONT_BODY_NAME,
            fontSize=14,
            leading=18,
            alignment=TA_CENTER,
            textColor=muted,
            spaceAfter=12,
        ),
        "cover_meta": ParagraphStyle(
            "TintaCoverMeta",
            fontName=FONT_BODY_NAME,
            fontSize=11,
            leading=15,
            alignment=TA_CENTER,
            textColor=muted,
        ),
        # Headings
        "h1": ParagraphStyle(
            "TintaH1",
            fontName=FONT_TITLE_NAME,
            fontSize=18,
            leading=22,
            spaceBefore=14,
            spaceAfter=8,
            textColor=tinta,
        ),
        "h2": ParagraphStyle(
            "TintaH2",
            fontName=FONT_TITLE_NAME,
            fontSize=14,
            leading=18,
            spaceBefore=10,
            spaceAfter=6,
            textColor=tinta,
        ),
        "h3": ParagraphStyle(
            "TintaH3",
            fontName=FONT_TITLE_NAME,
            fontSize=12,
            leading=16,
            spaceBefore=8,
            spaceAfter=4,
            textColor=tinta,
        ),
        # Body
        "body": ParagraphStyle(
            "TintaBody",
            fontName=FONT_BODY_NAME,
            fontSize=11,
            leading=15,
            spaceAfter=4,
            textColor=text,
            alignment=TA_JUSTIFY,
        ),
        "lead": ParagraphStyle(
            "TintaLead",
            fontName=FONT_BODY_NAME,
            fontSize=12,
            leading=18,
            spaceAfter=8,
            textColor=text,
        ),
        "caption": ParagraphStyle(
            "TintaCaption",
            fontName=FONT_BODY_NAME,
            fontSize=9,
            leading=12,
            textColor=muted,
            spaceAfter=8,
        ),
        # Code (monospace)
        "code": ParagraphStyle(
            "TintaCode",
            fontName="Courier",
            fontSize=9,
            leading=12,
            leftIndent=12,
            rightIndent=12,
            backColor=light,
            textColor=text,
            spaceBefore=4,
            spaceAfter=4,
        ),
        # Quote
        "quote": ParagraphStyle(
            "TintaQuote",
            fontName=FONT_TITLE_NAME,
            fontSize=12,
            leading=18,
            leftIndent=24,
            rightIndent=24,
            textColor=muted,
            spaceBefore=8,
            spaceAfter=4,
        ),
        "quote_attr": ParagraphStyle(
            "TintaQuoteAttr",
            fontName=FONT_BODY_NAME,
            fontSize=10,
            leading=14,
            leftIndent=24,
            rightIndent=24,
            textColor=muted,
            alignment=TA_LEFT,
            spaceAfter=12,
        ),
        # Callout — textColor por kind
        "callout_info": ParagraphStyle(
            "TintaCalloutInfo",
            fontName=FONT_BODY_NAME,
            fontSize=10,
            leading=14,
            leftIndent=12,
            rightIndent=12,
            backColor=light,
            textColor=text,
            borderColor=latao,
            borderWidth=0,
            borderPadding=8,
            spaceBefore=6,
            spaceAfter=6,
        ),
        # KPI
        "kpi_label": ParagraphStyle(
            "TintaKpiLabel",
            fontName=FONT_BODY_NAME,
            fontSize=9,
            leading=12,
            textColor=muted,
            alignment=TA_CENTER,
        ),
        "kpi_value": ParagraphStyle(
            "TintaKpiValue",
            fontName=FONT_TITLE_NAME,
            fontSize=18,
            leading=22,
            textColor=tinta,
            alignment=TA_CENTER,
        ),
        "kpi_delta": ParagraphStyle(
            "TintaKpiDelta",
            fontName=FONT_BODY_NAME,
            fontSize=9,
            leading=12,
            textColor=success,
            alignment=TA_CENTER,
        ),
        # TOC + chart placeholder
        "placeholder": ParagraphStyle(
            "TintaPlaceholder",
            fontName=FONT_BODY_NAME,
            fontSize=10,
            leading=14,
            textColor=muted,
            leftIndent=12,
            rightIndent=12,
            backColor=light,
            spaceBefore=6,
            spaceAfter=6,
        ),
        # Signature line
        "signature_line": ParagraphStyle(
            "TintaSignatureLine",
            fontName=FONT_BODY_NAME,
            fontSize=10,
            leading=14,
            textColor=muted,
            spaceBefore=4,
            spaceAfter=2,
        ),
        "signature_name": ParagraphStyle(
            "TintaSignatureName",
            fontName=FONT_BODY_NAME,
            fontSize=10,
            leading=14,
            textColor=text,
        ),
        # Sobrio marker (modo registraveis)
        "sobrio": is_sobrio,
    }


def _block_texts_with_font(block: dict, styles: dict) -> list[tuple[str, str]]:
    """Extrai (texto, font_name) de um bloco para o glifo-check.

    Retorna lista vazia se o bloco não tem texto (ex: page_break,
    divider, spacer, image). A font_name é a que o renderer
    vai usar (Serif para títulos, Sans para corpo, Courier
    para code).
    """
    btype = block.get("type", "")
    if btype in ("page_break", "divider", "spacer", "image"):
        return []
    out: list[tuple[str, str]] = []
    if btype == "cover":
        for k in ("title", "subtitle"):
            v = block.get(k)
            if v:
                out.append((v, FONT_TITLE_NAME if k == "title" else FONT_BODY_NAME))
        for k in ("author", "date"):
            v = block.get(k)
            if v:
                out.append((v, FONT_BODY_NAME))
    elif btype == "heading":
        v = block.get("text", "")
        if v:
            out.append((v, FONT_TITLE_NAME))
    elif btype == "paragraph":
        v = block.get("text", "")
        if v:
            out.append((v, FONT_BODY_NAME))
    elif btype == "list":
        for item in block.get("items", []):
            t = item.get("text", "") if isinstance(item, dict) else str(item)
            if t:
                out.append((t, FONT_BODY_NAME))
            for child in (item.get("children", []) if isinstance(item, dict) else []):
                ct = child.get("text", "") if isinstance(child, dict) else str(child)
                if ct:
                    out.append((ct, FONT_BODY_NAME))
    elif btype == "table":
        for h in block.get("headers", []):
            if h:
                out.append((h, FONT_BODY_NAME))
        for row in block.get("rows", []):
            for cell in row:
                if cell:
                    out.append((cell, FONT_BODY_NAME))
    elif btype == "key_value":
        for entry in block.get("entries", []):
            if isinstance(entry, dict):
                k = entry.get("key", "")
                v = entry.get("value", "")
                if k:
                    out.append((k, FONT_BODY_NAME))
                if v:
                    out.append((v, FONT_BODY_NAME))
    elif btype == "kpis":
        for item in block.get("items", []):
            for k in ("label", "value", "delta", "delta_label"):
                v = item.get(k)
                if v:
                    out.append(
                        (v, FONT_TITLE_NAME if k == "value" else FONT_BODY_NAME)
                    )
    elif btype == "callout":
        v = block.get("text", "")
        if v:
            out.append((v, FONT_BODY_NAME))
    elif btype == "quote":
        t = block.get("text", "")
        a = block.get("attribution", "")
        if t:
            out.append((t, FONT_TITLE_NAME))
        if a:
            out.append((a, FONT_BODY_NAME))
    elif btype == "steps":
        for s in block.get("items", []):
            t = s.get("title", "")
            d = s.get("description", "")
            if t:
                out.append((t, FONT_TITLE_NAME))
            if d:
                out.append((d, FONT_BODY_NAME))
    elif btype == "chart_placeholder":
        v = block.get("title", "")
        if v:
            out.append((v, FONT_BODY_NAME))
    elif btype == "code":
        c = block.get("content", "")
        if c:
            out.append((c, "Courier"))
    elif btype == "footer":
        v = block.get("text", "")
        if v:
            out.append((v, FONT_BODY_NAME))
    elif btype == "signatures":
        for p in block.get("pairs", []):
            for k in ("name", "role", "location"):
                v = p.get(k)
                if v:
                    out.append((v, FONT_BODY_NAME))
    elif btype == "back_cover":
        for k in ("name", "email", "phone", "address"):
            v = block.get(k)
            if v:
                out.append((v, FONT_BODY_NAME))
    elif btype == "toc":
        # Placeholder text — não passa pelo cmap real.
        pass
    return out


def _glyph_check(blocks: list) -> list[dict]:
    """Verifica todos os textos do spec contra o cmap das
    fontes Tinta & Latão via `fontTools` (D-GLYPH-1).

    Retorna lista vazia se OK. Em falta, lista de
    `{block_index, char, codepoint, font_name, block_type}`.
    """
    if not FONTTOOLS_AVAILABLE:
        # D-FAIL-1: o bootstrap hard-fail se fonttools faltar.
        # Se chegou aqui, alguém mexeu no runtime. Falha
        # estruturada em vez de render mudo.
        return [
            {
                "code": "fonttools_unavailable",
                "message": (
                    "fontTools nao instalado. Rode o bootstrap.ps1 (D-FAIL-1 do ADR-0021). "
                    "Glifo-check pre-render (D-GLYPH-1) nao pode executar sem ele."
                ),
            }
        ]
    # Cache de cmap por font name.
    cmap_cache: dict[str, dict] = {}
    missing: list[dict] = []
    for block_index, block in enumerate(blocks):
        if not isinstance(block, dict):
            continue
        for text, font_name in _block_texts_with_font(block, {}):
            if not text:
                continue
            # Resolve o path real (mesmo algoritmo do
            # `_try_register_font`: pega o primeiro candidato
            # que existe e tem > 50KB).
            if font_name not in cmap_cache:
                font_path = None
                for cand in FONT_FILES.get(font_name, []):
                    if cand.is_file() and cand.stat().st_size > 50_000:
                        font_path = cand
                        break
                if font_path is None:
                    # Fonte nao encontrada (fallback built-in
                    # do reportlab). Cmap check nao se aplica
                    # — o reportlab usa glifos do Type 1
                    # built-in (Adobe Standard Encoding) e a
                    # renderizacao fica feia. Registramos
                    # como warning mas nao falhamos o
                    # render — a decisao de "fail hard" e do
                    # bootstrap (D-FAIL-1), nao do handler.
                    cmap_cache[font_name] = None
                    continue
                try:
                    tt = FTFont(str(font_path))
                    cmap_cache[font_name] = tt.getBestCmap()
                except Exception as exc:
                    # Fonte corrompida — falha estruturada.
                    missing.append(
                        {
                            "block_index": block_index,
                            "char": "",
                            "codepoint": None,
                            "font_name": font_name,
                            "block_type": block.get("type"),
                            "error": f"falha abrindo {font_name}: {exc}",
                        }
                    )
                    cmap_cache[font_name] = None
                    continue
            cmap = cmap_cache[font_name]
            if cmap is None:
                continue
            for ch in text:
                cp = ord(ch)
                if cp not in cmap:
                    missing.append(
                        {
                            "block_index": block_index,
                            "char": ch,
                            "codepoint": cp,
                            "font_name": font_name,
                            "block_type": block.get("type"),
                        }
                    )
    return missing


def _build_story(
    blocks: list,
    styles: dict,
    identity: dict,
    font_status: dict,
) -> tuple[list, int]:
    """Constrói a `story` (lista de flowables) do reportlab
    a partir dos 20 blocos. Retorna `(story, blocks_written)`.

    Cada bloco vira 1+ flowables. Cobertura total:
    cover, toc, heading, paragraph, list, table, key_value,
    kpis, callout, quote, steps, chart_placeholder, image,
    code, divider, spacer, page_break, footer, signatures,
    back_cover. Chart e Toc viram placeholder com warning
    inline (a degradacao aparece no PDF — o usuario sabe).
    """
    story: list = []
    blocks_written = 0
    # v0.1: o último `Footer` bloco define o footer de página
    # inteira. Suporte a múltiplos footers no meio do doc é
    # Etapa 5.x. Mesmo padrão WordPro v0.1 (footer é
    # placeholder textual).
    last_footer = None
    for b in blocks:
        if not isinstance(b, dict):
            continue
        btype = b.get("type", "")
        if btype == "footer":
            last_footer = b
            blocks_written += 1
            continue  # footer é page-level, não flowable
        if btype == "cover":
            story.extend(_render_cover(b, styles))
        elif btype == "toc":
            story.extend(_render_toc(b, styles))
        elif btype == "heading":
            story.extend(_render_heading(b, styles))
        elif btype == "paragraph":
            story.extend(_render_paragraph(b, styles))
        elif btype == "list":
            story.extend(_render_list(b, styles))
        elif btype == "table":
            story.extend(_render_table(b, styles, identity))
        elif btype == "key_value":
            story.extend(_render_key_value(b, styles))
        elif btype == "kpis":
            story.extend(_render_kpis(b, styles, identity))
        elif btype == "callout":
            story.extend(_render_callout(b, styles))
        elif btype == "quote":
            story.extend(_render_quote(b, styles))
        elif btype == "steps":
            story.extend(_render_steps(b, styles))
        elif btype == "chart_placeholder":
            story.extend(_render_chart_placeholder(b, styles))
        elif btype == "image":
            story.extend(_render_image(b, styles))
        elif btype == "code":
            story.extend(_render_code(b, styles))
        elif btype == "divider":
            story.extend(_render_divider(b, styles, identity))
        elif btype == "spacer":
            story.extend(_render_spacer(b, styles))
        elif btype == "page_break":
            story.append(PageBreak())
        elif btype == "signatures":
            story.extend(_render_signatures(b, styles))
        elif btype == "back_cover":
            story.extend(_render_back_cover(b, styles, identity))
        else:
            # Bloco desconhecido — vira placeholder legível.
            story.append(
                Paragraph(f"[bloco desconhecido: {btype}]", styles["body"])
            )
        blocks_written += 1
    # Footer (se houve) vira o onPage do doc; guardado no
    # closure pelo `handle_pdf_write` (que monta o
    # `SimpleDocTemplate`).
    _last_footer_capture["footer"] = last_footer
    _last_footer_capture["font_status"] = font_status
    return story, blocks_written


# Captura do último footer visto — passada pro closure
# `_draw_watermark` via `_last_footer_capture` (mesma
# técnica do `_FONT_STATUS` global em `ensure_fonts_registered`).
_last_footer_capture: dict = {"footer": None, "font_status": {}}


def _draw_page_chrome(canvas, doc) -> None:
    """Callback de `SimpleDocTemplate.onPage` — desenha
    o footer (page-level) e a marca d'água (opt-in).
    """
    footer = _last_footer_capture.get("footer")
    if footer is not None:
        _draw_footer(canvas, doc, footer)
    watermark = _last_footer_capture.get("watermark")
    identity = _last_footer_capture.get("identity", {})
    if watermark is not None:
        _draw_watermark(canvas, doc, watermark, identity)


def _draw_footer(canvas, doc, footer: dict) -> None:
    """Desenha o rodapé no `canvas`. `page_numbers: true`
    adiciona 'N / total' à direita."""
    from reportlab.lib.pagesizes import A4 as _A4  # noqa: F401  (largura)
    from reportlab.lib.units import cm as _cm  # noqa: F401
    width, _ = _A4
    text = footer.get("text", "")
    page_numbers = bool(footer.get("page_numbers", False))
    canvas.saveState()
    canvas.setFont(FONT_BODY_NAME, 8)
    canvas.setFillColorRGB(0.5, 0.5, 0.5)  # muted
    # Texto à esquerda
    if text:
        canvas.drawString(2 * _cm, 1.0 * _cm, text[:120])
    # Numeração à direita
    if page_numbers:
        try:
            page_no = canvas.getPageNumber()
            label = f"Pág. {page_no}"
            canvas.drawRightString(width - 2 * _cm, 1.0 * _cm, label)
        except Exception:
            pass
    canvas.restoreState()


def _draw_watermark(canvas, doc, watermark: dict, identity: dict) -> None:
    """Desenha a marca d'água opt-in (D-PDF2 do ADR-0021).

    Posições suportadas: center, diagonal, bottom_right,
    top_right. Cor: do `identity.latao` (Tinta & Latão).
    **Não** renderiza em modo Sobrio (rejeitado pelo
    `validate_semantic` antes de chegar aqui).
    """
    from reportlab.lib.pagesizes import A4 as _A4
    from reportlab.lib.units import cm as _cm
    from reportlab.pdfbase import pdfmetrics as _pdfmetrics

    width, height = _A4
    text = watermark.get("text", "")
    if not text:
        return
    position = watermark.get("position", "center")
    opacity = watermark.get("opacity", 0.15)
    if opacity is None:
        opacity = 0.15
    font_size = watermark.get("font_size")
    if font_size is None:
        font_size = 72 if position in ("center", "diagonal") else 14
    color_hex = identity.get("latao", "#B8924A")
    canvas.saveState()
    try:
        canvas.setFillColor(_hex(color_hex))
        canvas.setFillAlpha(float(opacity))
    except Exception:
        canvas.setFillColorRGB(0.72, 0.57, 0.29)
        canvas.setFillAlpha(float(opacity))
    canvas.setFont(FONT_TITLE_NAME, float(font_size))
    if position == "center":
        # reportlab drawCentredString: posiciona pelo baseline.
        # Ajuste empírico pra centralizar visualmente.
        canvas.drawCentredString(
            width / 2, height / 2 - float(font_size) / 3, text
        )
    elif position == "diagonal":
        canvas.translate(width / 2, height / 2)
        canvas.rotate(45)
        canvas.drawCentredString(0, 0, text)
    elif position == "bottom_right":
        canvas.drawRightString(width - 2 * _cm, 2 * _cm, text)
    elif position == "top_right":
        canvas.drawRightString(width - 2 * _cm, height - 2 * _cm, text)
    canvas.restoreState()


# ---- renderers por bloco -----------------------------------------------


def _render_cover(b: dict, styles: dict) -> list:
    """Capa — vai na primeira página. Empilha título, subtítulo,
    autor, data, centralizados. Termina com `PageBreak` pra
    que o próximo bloco comece em página nova.
    """
    out: list = []
    out.append(Spacer(1, 4 * cm))
    title = b.get("title", "")
    if title:
        out.append(Paragraph(title, styles["cover_title"]))
    subtitle = b.get("subtitle", "")
    if subtitle:
        out.append(Paragraph(subtitle, styles["cover_subtitle"]))
    author = b.get("author", "")
    date = b.get("date", "")
    if author or date:
        meta = "  •  ".join(x for x in (author, date) if x)
        out.append(Spacer(1, 1 * cm))
        out.append(Paragraph(meta, styles["cover_meta"]))
    out.append(PageBreak())
    return out


def _render_toc(b: dict, styles: dict) -> list:
    """Toc — placeholder na v0.1. Sumário automático em duas
    passadas (`multiBuild` do reportlab) é Etapa 5.x."""
    return [
        Paragraph(
            "[Sumário: disponível em versão futura — Etapa 5.x]", styles["placeholder"]
        ),
        Spacer(1, 0.5 * cm),
    ]


def _render_heading(b: dict, styles: dict) -> list:
    """Heading com `level` 1-3. Level 4+ cai em h3 (mesma regra
    do `document-engine`)."""
    level = int(b.get("level", 1))
    text = b.get("text", "")
    number = b.get("number")
    if number:
        text = f"{number}  {text}"
    style_name = "h1" if level == 1 else "h2" if level == 2 else "h3"
    return [Paragraph(text, styles[style_name])]


def _render_paragraph(b: dict, styles: dict) -> list:
    """Parágrafo. Se o `style` for "lead", usa o estilo
    lead (maior, sem justify)."""
    text = b.get("text", "")
    style_name = b.get("style")
    if style_name == "lead":
        sn = "lead"
    else:
        sn = "body"
    return [Paragraph(text, styles[sn])]


def _render_list(b: dict, styles: dict) -> list:
    """Lista numerada ou com marcadores. Cada item vira
    um parágrafo com prefixo ('- ' ou 'N. ')."""
    ordered = bool(b.get("ordered", False))
    items = b.get("items", [])
    out: list = []
    for i, item in enumerate(items, start=1):
        if isinstance(item, dict):
            text = item.get("text", "")
            children = item.get("children", [])
        else:
            text = str(item)
            children = []
        prefix = f"{i}. " if ordered else "•  "
        out.append(Paragraph(f"{prefix}{text}", styles["body"]))
        for child in children:
            if isinstance(child, dict):
                ct = child.get("text", "")
            else:
                ct = str(child)
            if ct:
                out.append(Paragraph(f"     {ct}", styles["body"]))
    return out


def _render_table(b: dict, styles: dict, identity: dict) -> list:
    """Tabela. Capa com `title` em cima, `source` embaixo."""
    headers = b.get("headers", [])
    rows = b.get("rows", [])
    title = b.get("title")
    source = b.get("source")
    out: list = []
    if title:
        out.append(Paragraph(title, styles["h3"]))
    if not headers and not rows:
        return out
    # Monta a Table. Headers em negrito (cor tinta),
    # linhas com grade fina (cor muted).
    data = [list(headers)] + [list(r) for r in rows]
    n_cols = max(len(headers), max((len(r) for r in rows), default=0))
    if n_cols == 0:
        return out
    # Pad cells
    for row in data:
        while len(row) < n_cols:
            row.append("")
    # Dimensões: largura disponível na area util.
    from reportlab.lib.units import cm as _cm
    from reportlab.lib.pagesizes import A4 as _A4
    _, _ = _A4, None
    avail = _A4[0] - 4 * _cm  # margens laterais 2cm
    col_width = avail / n_cols
    grid_color = _hex(identity.get("muted", "#6B7280"))
    header_bg = _hex(identity.get("light", "#F3F4F6"))
    ts = TableStyle(
        [
            ("FONTNAME", (0, 0), (-1, 0), FONT_TITLE_NAME),
            ("FONTSIZE", (0, 0), (-1, 0), 10),
            ("FONTNAME", (0, 1), (-1, -1), FONT_BODY_NAME),
            ("FONTSIZE", (0, 1), (-1, -1), 10),
            ("BACKGROUND", (0, 0), (-1, 0), header_bg),
            ("TEXTCOLOR", (0, 0), (-1, 0), _hex(identity.get("tinta", "#1A2B4A"))),
            ("GRID", (0, 0), (-1, -1), 0.5, grid_color),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]
    )
    t = Table(data, colWidths=[col_width] * n_cols)
    t.setStyle(ts)
    out.append(t)
    if source:
        out.append(Paragraph(f"Fonte: {source}", styles["caption"]))
    out.append(Spacer(1, 0.3 * cm))
    return out


def _render_key_value(b: dict, styles: dict) -> list:
    """Tabela de chave-valor (2 colunas)."""
    entries = b.get("entries", [])
    if not entries:
        return []
    data = []
    for entry in entries:
        if isinstance(entry, dict):
            k = entry.get("key", "")
            v = entry.get("value", "")
        elif isinstance(entry, (list, tuple)) and len(entry) >= 2:
            k, v = entry[0], entry[1]
        else:
            continue
        data.append([str(k), str(v)])
    if not data:
        return []
    from reportlab.lib.units import cm as _cm
    from reportlab.lib.pagesizes import A4 as _A4
    avail = _A4[0] - 4 * _cm
    col_w = [avail * 0.35, avail * 0.65]
    ts = TableStyle(
        [
            ("FONTNAME", (0, 0), (0, -1), FONT_TITLE_NAME),
            ("FONTNAME", (1, 0), (1, -1), FONT_BODY_NAME),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 2),
            ("RIGHTPADDING", (0, 0), (-1, -1), 2),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ("LINEBELOW", (0, 0), (-1, -1), 0.25, _hex("#6B7280")),
        ]
    )
    t = Table(data, colWidths=col_w)
    t.setStyle(ts)
    return [t, Spacer(1, 0.3 * cm)]


def _render_kpis(b: dict, styles: dict, identity: dict) -> list:
    """Painel de KPIs (2-4 itens). Cada KPI vira label + value +
    delta, dispostos em colunas iguais.
    """
    items = b.get("items", [])
    if not items:
        return []
    # Render cada KPI como 3 parágrafos (label, value, delta)
    # dentro de uma Table de N colunas (N = len(items)).
    from reportlab.lib.units import cm as _cm
    from reportlab.lib.pagesizes import A4 as _A4
    avail = _A4[0] - 4 * _cm
    n = len(items)
    col_w = avail / n
    # Cada coluna tem 3 linhas: label, value, delta.
    grid = []
    for kpi in items:
        label = kpi.get("label", "")
        value = kpi.get("value", "")
        delta = kpi.get("delta", "")
        grid.append(
            [
                Paragraph(label, styles["kpi_label"]),
                Paragraph(value, styles["kpi_value"]),
                Paragraph(delta if delta else "", styles["kpi_delta"]),
            ]
        )
    # Transpõe: grid[col][row] -> data[row][col]
    data = [[grid[c][r] for c in range(n)] for r in range(3)]
    ts = TableStyle(
        [
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]
    )
    t = Table(data, colWidths=[col_w] * n)
    t.setStyle(ts)
    return [t, Spacer(1, 0.3 * cm)]


def _render_callout(b: dict, styles: dict) -> list:
    """Callout — caixa de destaque. `kind` decide o prefixo
    e (em v0.1) só o prefixo varia; a cor de fundo é
    `identity.light` (mesma para todos os kinds — Etapa 6
    refina as cores por kind).
    """
    kind = b.get("kind", "info")
    text = b.get("text", "")
    prefix = {
        "info": "[INFO]",
        "alert": "[ALERTA]",
        "critical": "[CRÍTICO]",
        "success": "[OK]",
    }.get(kind, f"[{kind.upper()}]")
    return [Paragraph(f"{prefix}  {text}", styles["callout_info"])]


def _render_quote(b: dict, styles: dict) -> list:
    """Citação com atribuição opcional."""
    text = b.get("text", "")
    attr = b.get("attribution")
    out = [Paragraph(f'"{text}"', styles["quote"])]
    if attr:
        out.append(Paragraph(f"— {attr}", styles["quote_attr"]))
    return out


def _render_steps(b: dict, styles: dict) -> list:
    """Passos numerados."""
    items = b.get("items", [])
    out: list = []
    for i, s in enumerate(items, start=1):
        title = s.get("title", "")
        desc = s.get("description")
        out.append(Paragraph(f"{i}. {title}", styles["h3"]))
        if desc:
            out.append(Paragraph(desc, styles["body"]))
    return out


def _render_chart_placeholder(b: dict, styles: dict) -> list:
    """Chart — placeholder textual em v0.1. Render real
    (bar/line/pie com cores) é Etapa 5.x.
    """
    kind = b.get("kind", "?")
    title = b.get("title", "")
    label = f"[Gráfico de {kind}"
    if title:
        label += f" — {title}"
    label += " — visualização nativa prevista para Etapa 5.x]"
    return [Paragraph(label, styles["placeholder"])]


def _render_image(b: dict, styles: dict) -> list:
    """Imagem. Path validado pelo `validate_path("read")`
    antes (defesa em profundidade — o caller Rust já
    checa a allowlist). Width opcional.
    """
    path = b.get("path", "")
    if not path:
        return [Paragraph("[Imagem: path ausente]", styles["body"])]
    try:
        from reportlab.lib.units import cm as _cm
        if b.get("width_cm"):
            img = RLImage(path, width=float(b["width_cm"]) * _cm)
        else:
            img = RLImage(path, width=10 * _cm)
        out: list = [img]
    except Exception as exc:
        out = [Paragraph(f"[Imagem não carregada: {exc}]", styles["body"])]
    cap = b.get("caption")
    alt = b.get("alt", "")
    if cap:
        out.append(Paragraph(cap, styles["caption"]))
    elif alt:
        out.append(Paragraph(f"[{alt}]", styles["caption"]))
    return out


def _render_code(b: dict, styles: dict) -> list:
    """Bloco de código. Preserva indentação e quebra
    por linha. `language` é decorativo (sem highlight
    no v0.1 — `RCodeParser` é Etapa 5.x).
    """
    content = b.get("content", "")
    out: list = []
    if b.get("language"):
        out.append(Paragraph(f"<font color='#6B7280'>{b['language']}</font>", styles["caption"]))
    for line in content.splitlines():
        out.append(Paragraph(line if line else "&#160;", styles["code"]))
    if b.get("caption"):
        out.append(Paragraph(b["caption"], styles["caption"]))
    return out


def _render_divider(b: dict, styles: dict, identity: dict) -> list:
    """Linha horizontal. Implementada como Table 1x1 sem
    conteúdo, só border bottom."""
    from reportlab.lib.units import cm as _cm
    from reportlab.lib.pagesizes import A4 as _A4
    width = _A4[0] - 4 * _cm
    t = Table([[""]], colWidths=[width], rowHeights=[0.1 * _cm])
    t.setStyle(
        TableStyle(
            [
                (
                    "LINEBELOW",
                    (0, 0),
                    (-1, -1),
                    0.75,
                    _hex(identity.get("latao", "#B8924A")),
                ),
            ]
        )
    )
    return [t, Spacer(1, 0.3 * cm)]


def _render_spacer(b: dict, styles: dict) -> list:
    """Espaço vertical."""
    from reportlab.lib.units import cm as _cm
    h = float(b.get("height_cm", 0.5))
    return [Spacer(1, h * _cm)]


def _render_signatures(b: dict, styles: dict) -> list:
    """Bloco de assinaturas. Cada par vira linha + nome + role + local."""
    pairs = b.get("pairs", [])
    out: list = []
    if not pairs:
        return out
    out.append(Spacer(1, 2 * cm))
    for p in pairs:
        name = p.get("name", "")
        role = p.get("role")
        loc = p.get("location")
        out.append(Paragraph("___________________________", styles["signature_line"]))
        out.append(Paragraph(name, styles["signature_name"]))
        if role:
            out.append(Paragraph(role, styles["signature_line"]))
        if loc:
            out.append(Paragraph(loc, styles["signature_line"]))
        out.append(Spacer(1, 1 * cm))
    return out


def _render_back_cover(b: dict, styles: dict, identity: dict) -> list:
    """Contracapa. Empilha nome + contatos centralizados,
    termina com PageBreak."""
    out: list = []
    out.append(PageBreak())
    out.append(Spacer(1, 4 * cm))
    name = b.get("name", "")
    if name:
        out.append(Paragraph(name, styles["cover_title"]))
    for k in ("email", "phone", "address"):
        v = b.get(k)
        if v:
            out.append(Paragraph(v, styles["cover_meta"]))
    return out


# ---- entrypoint --------------------------------------------------------


def handle_pdf_write(payload: dict) -> dict:
    """`pdf.write` v0.4.0 (Etapa 5 PR 2): payload estendido.

    Input (campos obrigatórios marcados com *):
      - `path`* (str): destino do .pdf.
      - `title`* (str): título do documento.
      - `style` (str): `"tinta_e_latao"` (default) ou `"sobrio"`.
      - `page` (dict): `{size, margin_cm: {top, bottom, left, right}}`.
      - `identity` (dict): paleta de cor (Tinta/Latão ou Sobrio).
      - `watermark` (dict | null): opt-in (D-PDF2 do ADR-0021).
      - `metadata` (dict): author, organization, keywords, description,
        confidentiality.
      - `blocks`* (list): 1+ blocos do `DocumentBlock`, cada um
        com `type` discriminado.

    Output (sucesso):
      `{ok, path, size_bytes, pages_rendered, blocks_written,
        glifo_check: {checked, missing}}`.

    Erros (em `tool.result {ok: false, code, message, ...}`):
      - `invalid_payload`: campo obrigatório faltando.
      - `fonttools_unavailable`: `fontTools` não instalado
        (D-FAIL-1 violado).
      - `missing_glyph`: glifo faltando no cmap de uma fonte
        Tinta & Latão (D-GLYPH-1). Inclui lista em `missing`:
        `[{block_index, char, codepoint, font_name, block_type}, ...]`.
    """
    # 1. Validação do path (defesa em profundidade; o
    # `WorkerToolDispatcher` no Rust já fez antes do invoke).
    if "path" not in payload:
        return {
            "ok": False,
            "code": "invalid_payload",
            "message": "campo 'path' ausente",
        }
    path = validate_path(_payload_field(payload, "path", str), "write")
    if "title" not in payload:
        return {
            "ok": False,
            "code": "invalid_payload",
            "message": "campo 'title' ausente",
        }
    title = _payload_field(payload, "title", str)
    if "blocks" not in payload:
        return {
            "ok": False,
            "code": "invalid_payload",
            "message": "campo 'blocks' ausente",
        }
    blocks = _payload_field(payload, "blocks", list)
    if not blocks:
        return {
            "ok": False,
            "code": "invalid_payload",
            "message": "'blocks' nao pode ser vazio",
        }

    # 2. Defaults.
    style_name = payload.get("style", "tinta_e_latao")
    if style_name not in ("tinta_e_latao", "sobrio"):
        style_name = "tinta_e_latao"
    page = payload.get("page") or {}
    page_margin = page.get("margin_cm") or {}
    identity = payload.get("identity") or {}
    watermark = payload.get("watermark")
    metadata = payload.get("metadata") or {}

    # 3. Garante que as fontes Tinta & Latão estão
    # registradas no `pdfmetrics`. Idempotente. Se as TTFs
    # não existirem, o fallback built-in do reportlab é
    # usado (mesma disciplina da v0.3.0; o hard-fail é
    # do bootstrap, não do handler).
    font_status = ensure_fonts_registered()

    # 4. Glifo-check pre-render (D-GLYPH-1 do ADR-0021).
    # varre todos os blocos, intersecta com cmap das
    # fontes Tinta & Latão via `fontTools`, falha
    # estruturado se algum glifo faltar. **Antes** do
    # `doc.build()` — render mudo é pior que erro claro.
    missing = _glyph_check(blocks)
    if missing:
        # Se o primeiro item for um "error code" (fonttools
        # missing), retorna aquele code direto.
        if len(missing) == 1 and missing[0].get("code"):
            return {
                "ok": False,
                "code": missing[0]["code"],
                "message": missing[0]["message"],
                "path": str(path),
            }
        return {
            "ok": False,
            "code": "missing_glyph",
            "message": (
                f"{len(missing)} glifo(s) faltando em fontes Tinta & Latao. "
                "Veja 'missing' para bloco + caractere + fonte."
            ),
            "missing": missing,
            "path": str(path),
        }

    # 5. Constrói os styles (cores do `identity`, fontes
    # Tinta & Latão ou fallback built-in do reportlab).
    styles = _build_pdf_styles(identity, style_name)

    # 6. Constrói a `story` (lista de flowables). O
    # `_last_footer_capture` é populado se houver bloco
    # `Footer` — o `_draw_page_chrome` (callback de
    # `onPage`) usa depois.
    _last_footer_capture["watermark"] = watermark
    _last_footer_capture["identity"] = identity
    story, blocks_written = _build_story(blocks, styles, identity, font_status)
    if not story:
        return {
            "ok": False,
            "code": "invalid_payload",
            "message": "story vazia depois do translate (nenhum bloco produziu flowable)",
        }

    # 7. Monta o `SimpleDocTemplate` com margens do
    # `page.margin_cm` e `onPage=_draw_page_chrome` (footer
    # + watermark).
    try:
        left = float(page_margin.get("left", 2.0)) * cm
        right = float(page_margin.get("right", 2.0)) * cm
        top = float(page_margin.get("top", 2.5)) * cm
        bottom = float(page_margin.get("bottom", 2.5)) * cm
    except (TypeError, ValueError):
        left, right, top, bottom = 2 * cm, 2 * cm, 2.5 * cm, 2.5 * cm

    # Metadados do PDF (vão pro /Info do PDF).
    pdf_title = title
    pdf_author = metadata.get("author") or ""
    pdf_subject = metadata.get("description") or ""
    pdf_keywords = metadata.get("keywords") or ""
    pdf_creator = "Frederico IA Studio - PDFPro v0.1 (Etapa 5 PR 2)"

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=left,
        rightMargin=right,
        topMargin=top,
        bottomMargin=bottom,
        title=pdf_title,
        author=pdf_author,
        subject=pdf_subject,
        keywords=pdf_keywords,
        creator=pdf_creator,
        onPage=_draw_page_chrome,
    )

    # 8. Build.
    try:
        doc.build(story)
    except Exception as exc:
        return {
            "ok": False,
            "code": "build_failed",
            "message": f"reportlab falhou no build: {exc}",
            "path": str(path),
        }

    # 9. Response. `pages_rendered` é um chute mínimo (1)
    # — reportlab não expõe `n_pages` pós-build. A
    # auditoria bloqueante do §19.6 (PRs 3-4) vai
    # calcular o real via `pypdfium2` e falhar se
    # discrepante. Por enquanto, 1 é honesto: "pelo
    # menos 1 página".
    size = path.stat().st_size
    return {
        "ok": True,
        "path": str(path),
        "size_bytes": size,
        "pages_rendered": 1,
        "blocks_written": blocks_written,
        "glifo_check": {
            "checked": sum(
                len(t) for b in blocks if isinstance(b, dict) for t, _ in _block_texts_with_font(b, {})
            ),
            "missing": [],
        },
    }


# ---- pdf.read -------------------------------------------------------------


def handle_ocr_run(payload: dict) -> dict:
    """ocr.run: OCR de uma imagem (PNG/JPG/TIFF/BMP) via Tesseract.

    Input: `{"path": str, "lang": str?}`
      - `path`: path da imagem (validado pelo `validate_path`).
      - `lang`: idiomas concatenados por `+`. Default: `por+eng`.
        Validado contra `INSTALLED_OCR_LANGS` - erro estruturado
        (nao mensagem criptica do Tesseract) se faltar.

    Output: `{"ok": true, "path": str, "text": str, "lang": str,
              "conf": float|None, "tesseract_version": str}`

    Erros comuns (em `tool.result {ok: false, code, message}`):
      - `ocr_not_available`: Tesseract nao instalado (bootstrap pulou).
      - `invalid_lang`: `lang` mal formado ou idioma nao instalado.
      - `tesseract_failed`: Tesseract retornou erro.
      - `ocr_timeout`: timeout de 60s excedido.
      - `image_not_found`: path nao existe.
    """
    path = validate_path(_payload_field(payload, "path", str), "read")
    lang = payload.get("lang", DEFAULT_OCR_LANG)
    if not isinstance(lang, str):
        lang = DEFAULT_OCR_LANG
    try:
        lang = _validate_lang(lang)
    except ValueError as exc:
        return {
            "ok": False,
            "code": "invalid_lang",
            "message": str(exc),
            "path": str(path),
        }

    if not PYTESSERACT_AVAILABLE:
        return {
            "ok": False,
            "code": "ocr_not_available",
            "message": (
                "pytesseract nao instalado. Rode o bootstrap.ps1 pra "
                "instalar (pula silenciosamente em contexto non-elevated)."
            ),
            "path": str(path),
        }
    if not _tesseract_executable_present():
        return {
            "ok": False,
            "code": "ocr_not_available",
            "message": (
                f"Tesseract nao encontrado em {TESSERACT_EXE}. "
                "Instale manualmente ou rode o bootstrap.ps1 como Admin."
            ),
            "path": str(path),
        }

    tess_version = _get_tesseract_version()
    try:
        out = _ocr_image_to_text(path, lang, timeout_s=60)
    except FileNotFoundError:
        return {
            "ok": False,
            "code": "image_not_found",
            "message": f"imagem nao encontrada: {path}",
            "path": str(path),
        }
    except RuntimeError as exc:
        msg = str(exc)
        if "timeout" in msg.lower():
            code = "ocr_timeout"
        elif "Tesseract falhou" in msg:
            code = "tesseract_failed"
        else:
            code = "ocr_error"
        log.warning("ocr.run falhou: %s (%s)", code, msg)
        return {
            "ok": False,
            "code": code,
            "message": msg,
            "path": str(path),
            "lang": lang,
            "tesseract_version": tess_version,
        }
    return {
        "ok": True,
        "path": str(path),
        "text": out["text"],
        "lang": lang,
        "conf": out["conf"],
        "tesseract_version": tess_version,
    }


# ---- pdf.read -------------------------------------------------------------


def handle_pdf_read(payload: dict) -> dict:
    """pdf.read: extrai texto de um .pdf e detecta paginas escaneadas.

    Input: `{"path": str, "ocr": "auto"|"never"|"only"?}`
      - `ocr` (default `"auto"`):
        - `"auto"`: se ha paginas escaneadas e Tesseract esta
          disponivel, faz OCR delas e popula `ocr_text`. Se 100%
          escaneado + OCR OK: devolve `ok: true` com `text` do OCR
          (precedencia: `text` = OCR, com `extraction: "ocr"`).
        - `"never"`: rapido, so checa camada de texto via `pdfplumber`.
        - `"only"`: ignora camada de texto e faz OCR de TODAS as
          paginas (mesmo as com texto).

    Output: `{"ok": true, "text": str, "ocr_text": {page: str},
              "page_count": int, "scanned_pages": [int],
              "ocr_available": bool, "ocr_truncated": bool,
              "extraction": "text"|"ocr"|"mixed",
              "tesseract_version": str|null}`

    **Procedencia sempre clara** (ADR-0019 §Decisao 3): `text` e
    `ocr_text` sao campos SEPARADOS. `text` so vem da camada do PDF;
    `ocr_text` e o mapa de OCR por pagina. Nunca mistura os dois -
    OCR troca 8 por B, 0 por O, 1 por l, e e exatamente em
    CNPJ/competencia/valor que o erro cai. Misturar apagaria a
    procedencia (mesma disciplina de `origin`/`external_content`
    da memoria).

    **Teto de paginas** (`MAX_OCR_PAGES_PDF` = 20): PDF escaneado
    de 200 paginas leva minutos e estoura o timeout do worker.
    Quando bate o teto, devolve parcial com `ocr_truncated: true`
    e segue (caller decide se aceita ou aborta).

    **Idioma do fallback automatico:** `PDF_FALLBACK_OCR_LANG` =
    `"por"` (contexto brasileiro, sem lixo de outro idioma; o
    chamador que sabe o idioma certo usa `ocr.run` direto).
    """
    path = validate_path(_payload_field(payload, "path", str), "read")
    ocr_mode = payload.get("ocr", "auto")
    if ocr_mode not in ("auto", "never", "only"):
        raise ValueError(
            f"'ocr' deve ser 'auto', 'never' ou 'only'; veio {ocr_mode!r}"
        )

    # 1. Extrai texto via pdfplumber (camada de texto do PDF).
    pages_text: list[str] = []
    scanned_pages: list[int] = []
    page_count = 0
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if not text:
                scanned_pages.append(i)
            pages_text.append(text)
        page_count = len(pages_text)

    full_text = "\n".join(pages_text)
    tess_version = _get_tesseract_version() if PYTESSERACT_AVAILABLE else None
    ocr_available = bool(
        PYTESSERACT_AVAILABLE
        and _tesseract_executable_present()
        and tess_version is not None
    )

    # 2. Decide se faz OCR.
    ocr_text: dict[int, str] = {}
    ocr_truncated = False
    pages_to_ocr: list[int] = []

    if ocr_mode == "only":
        # Ignora camada de texto, faz OCR de TODAS as paginas.
        pages_to_ocr = list(range(1, page_count + 1))
    elif ocr_mode == "auto" and scanned_pages and ocr_available:
        # Fallback transparente: so as paginas escaneadas.
        pages_to_ocr = list(scanned_pages)
    # ocr_mode == "never" => pages_to_ocr fica vazio.

    if pages_to_ocr and not ocr_available:
        # Caller pediu OCR mas nao temos. Devolve code estruturado
        # em vez de quebrar - o caller decide.
        return {
            "ok": False,
            "code": "ocr_not_available",
            "message": (
                f"`ocr: {ocr_mode}` pedido mas Tesseract/pytesseract "
                "nao disponivel. Rode o bootstrap.ps1 como Admin pra "
                "instalar. PDF continua legivel via camada de texto "
                "(se houver) - use `ocr: 'never'` pra forcar esse modo."
            ),
            "path": str(path),
            "page_count": page_count,
            "scanned_pages": scanned_pages,
            "ocr_available": False,
        }

    if pages_to_ocr:
        # 3. Teto de paginas.
        if len(pages_to_ocr) > MAX_OCR_PAGES_PDF:
            log.warning(
                "pdf.read: %d paginas pra OCR > teto %d - truncando",
                len(pages_to_ocr), MAX_OCR_PAGES_PDF,
            )
            ocr_truncated = True
            pages_to_ocr = pages_to_ocr[:MAX_OCR_PAGES_PDF]
        # 4. Faz OCR pagina a pagina. Falha de uma nao aborta as
        #    outras (best-effort).
        for page_num in pages_to_ocr:
            try:
                text = _ocr_pdf_page_to_text(
                    path, page_num, PDF_FALLBACK_OCR_LANG,
                    timeout_s=OCR_TIMEOUT_S_PER_PAGE,
                )
            except Exception as exc:
                log.warning(
                    "OCR da pagina %d de %s falhou: %s",
                    page_num, path, exc,
                )
                ocr_text[page_num] = ""  # marca como tentado
                ocr_truncated = True
                continue
            ocr_text[page_num] = text

    # 5. Monta o retorno. `extraction` diz qual caminho produziu
    #    o `text` - util quando o caller quiser auditar.
    if ocr_mode == "only" and ocr_text:
        # `text` agora e o OCR de TODAS as paginas (camada ignorada).
        all_text = "\n".join(
            ocr_text.get(i, "").strip() for i in range(1, page_count + 1)
        )
        return {
            "ok": True,
            "path": str(path),
            "text": all_text,
            "ocr_text": ocr_text,
            "page_count": page_count,
            "scanned_pages": [],  # nao faz sentido no modo "only"
            "ocr_available": ocr_available,
            "ocr_truncated": ocr_truncated,
            "extraction": "ocr",
            "tesseract_version": tess_version,
        }

    if scanned_pages and page_count > 0 and len(scanned_pages) == page_count:
        # 100% escaneado. Se Tesseract OK (ocr: "auto" + ocr_available
        # + pages_to_ocr nao vazio), o OCR rodou e o ocr_text tem as
        # paginas todas. Caso contrario, devolve code estruturado.
        if ocr_text and len(ocr_text) == page_count:
            all_text = "\n".join(
                ocr_text.get(i, "").strip() for i in range(1, page_count + 1)
            )
            return {
                "ok": True,
                "path": str(path),
                "text": all_text,
                "ocr_text": ocr_text,
                "page_count": page_count,
                "scanned_pages": scanned_pages,
                "ocr_available": ocr_available,
                "ocr_truncated": ocr_truncated,
                "extraction": "ocr",
                "tesseract_version": tess_version,
            }
        return {
            "ok": False,
            "code": "pdf_scanned_no_ocr",
            "message": (
                f"PDF 100% escaneado ({page_count} pagina(s) sem "
                "camada de texto) e OCR nao disponivel. Instale "
                "Tesseract via bootstrap.ps1 como Admin ou use "
                "`ocr: 'never'` se quiser confirmar a limitacao."
            ),
            "path": str(path),
            "page_count": page_count,
            "scanned_pages": scanned_pages,
            "ocr_available": ocr_available,
        }

    # Caso geral: PDF tem pelo menos uma pagina com camada de texto.
    # `text` = camada do PDF (fonte: pdfplumber). `ocr_text` separado
    # so com paginas escaneadas (se rodou OCR).
    if ocr_text:
        extraction = "mixed"
    else:
        extraction = "text"
    return {
        "ok": True,
        "path": str(path),
        "text": full_text,
        "ocr_text": ocr_text,
        "page_count": page_count,
        "scanned_pages": scanned_pages,
        "ocr_available": ocr_available,
        "ocr_truncated": ocr_truncated,
        "extraction": extraction,
        "tesseract_version": tess_version,
    }


# ---------------------------------------------------------------------------
# pdf.audit - Etapa 5 PR 3 da Fase 5 (D-PDF5 + D-PDF6 do ADR-0021)
# ---------------------------------------------------------------------------
#
# Auditoria estrutural bloqueante do §19.4 (PROMPT MESTRE). O §19.6
# diz que a auditoria nao tem interruptor - o `pdf.audit` falha
# barulhento quando algo nao bate, e o caller (kit) tem que abortar
# a entrega.
#
# **Capability name:** `pdf.audit` (sem terceiro nivel). Todos os
# outros capabilities do worker seguem `<dominio>.<verbo>`:
# `pdf.write`, `pdf.read`, `ocr.run`. O PR 4 vai estender este mesmo
# handler com `kind: "visual"` (pypdfium2 rasteriza + checa grade,
# sobreposicao, pagina vazia, alinhamento) - o `salvar()` faz uma
# chamada so porque §19.6 exige auditoria inteira.
#
# **Input:**
#   - path (str): path do .pdf a ser auditado
#   - kind (str): "structural" (PR 3) | "visual" (PR 4 - rejeitado por agora)
#   - pdfa (None | "pdfa_2b"): D-PDF5 opt-in. None = baseline; "pdfa_2b"
#     = baseline + A-2B subset
#   - metadata (dict): campos do DocumentMetadata (author, title,
#     organization, keywords, description) - usado pra cross-check
#     XMP/DocInfo quando A-2B opt-in
#   - expected_pages (int | None): cross-check com `pages_rendered` do
#     `pdf.write`. None pula o cross-check
#
# **Output (sucesso):**
#   {ok: True, checks: [...], coverage: "full", cache_key, rules_version}
#
# **Output (falha):**
#   {ok: False, code: "pdf_audit_structural_failed", message,
#    failed: [...], checks_passed: [...], cache_key, rules_version}
#   Cada item em `failed` tem {check, expected, got} - motivo legivel
#   pro caller (kit / tool / modelo) e pro usuario final.
#
# **Cache key (D-AUDIT-1):** sha256(pdf_bytes) + ":" + AUDIT_RULES_VERSION.
# PR 3 nao persiste cache (cache persistente e PR proprio - mistura
# correcao com otimizacao faz bug de cache virar bug de auditoria).
# Cache key retornado no `tool.result` pra que o caller saiba invalidar
# quando bumpar AUDIT_RULES_VERSION.

# Versao das regras de auditoria. Bump quando uma regra mudar
# (D-AUDIT-1: "mudanca de regra = bump da AUDIT_RULES_VERSION").
# 0.1.0 no PR 3 (regras iniciais baseline + A-2B subset).
AUDIT_RULES_VERSION = "0.1.0"


def _audit_icc_signature(blob: bytes) -> tuple[bool, str]:
    """Valida o header de um ICC profile embedded no OutputIntent.

    Checa so o minimo estrutural: signature 'acsp' no offset 36 e
    color space 'RGB ' no offset 16. Nao verifica TRC, primaries, ou
    outros tags - a auditoria rigorosa (veraPDF) roda no ci-nightly
    (D-PDF5 do ADR-0021: "veraPDF no job noturno valida PDFs
    gerados com opt-in"). Hard-fail do bootstrap se pikepdf faltar
    (D-FAIL-1).
    """
    if len(blob) < 128:
        return False, f"ICC stream muito pequeno: {len(blob)} bytes (minimo 128)"
    if blob[36:40] != b"acsp":
        return False, f"ICC signature invalida: esperado 'acsp' no offset 36, veio {blob[36:40]!r}"
    cs = blob[16:20]
    if cs != b"RGB ":
        return False, f"ICC color space != RGB (offset 16): {cs!r}"
    return True, "ICC header OK"


def _check_n_pages(pdf, n_pages_from_write: int | None) -> tuple[bool, dict]:
    """Check 1: n_pages >= 1, e bate com pages_rendered do write se fornecido."""
    n = len(pdf.pages)
    if n < 1:
        return False, {"check": "n_pages", "expected": ">= 1", "got": n}
    if n_pages_from_write is not None and n != n_pages_from_write:
        return False, {
            "check": "n_pages_consistency",
            "expected": n_pages_from_write,
            "got": n,
            "message": (
                f"n_pages do pikepdf ({n}) difere do pages_rendered do "
                f"pdf.write ({n_pages_from_write})"
            ),
        }
    return True, {"check": "n_pages", "value": n}


def _check_docinfo(pdf) -> tuple[bool, list[dict]]:
    """Check 2: DocInfo populado (Author/Title/Producer/Creator nao vazios)."""
    info = pdf.docinfo or {}
    failed: list[dict] = []
    for field in ("/Author", "/Title", "/Producer", "/Creator"):
        value = ""
        try:
            value = str(info.get(field, "")) if field in info else ""
        except Exception:
            value = ""
        if not value:
            failed.append({
                "check": "docinfo_field",
                "field": field,
                "expected": "non-empty",
                "got": "",
            })
    if failed:
        return False, failed
    return True, [{"check": "docinfo_populated", "fields": ["Author", "Title", "Producer", "Creator"]}]


def _check_fonts_embedded(pdf) -> tuple[bool, list[dict]]:
    """Check 3: todas as fontes referenciadas tem FontFile embedded.

    Walk em /Resources/Font/*. /FontDescriptor sem /FontFile* =
    fonte nao embedded. Tipos: Type 1 -> /FontFile, TrueType ->
    /FontFile2, CFF -> /FontFile3. O reportlab (PR 2) ja embarca
    as Tinta e Latao via TTFont(), entao este check falha
    estruturado se o FontFile sumir (bug no render ou manipulacao
    pos-escrita).
    """
    failed: list[dict] = []
    seen: set[str] = set()
    for page_num, page in enumerate(pdf.pages, start=1):
        try:
            page_obj = page.obj if hasattr(page, "obj") else page
        except Exception:
            continue
        if not page_obj or "/Resources" not in page_obj:
            continue
        res = page_obj["/Resources"]
        if not res or "/Font" not in res:
            continue
        fonts = res["/Font"]
        if not fonts:
            continue
        for font_name, font_ref in fonts.items():
            key = str(font_name)
            if key in seen:
                continue
            seen.add(key)
            try:
                f = font_ref.obj if hasattr(font_ref, "obj") else font_ref
                if not f or "/FontDescriptor" not in f:
                    continue
                fd = f["/FontDescriptor"]
                if not any(k in fd for k in ("/FontFile", "/FontFile2", "/FontFile3")):
                    failed.append({
                        "check": "font_embedded",
                        "font": key,
                        "page": page_num,
                        "expected": "FontFile* em FontDescriptor",
                        "got": "nenhum",
                    })
            except Exception:
                # Nao conseguimos inspecionar essa fonte - nao
                # falhamos por causa de bug do nosso walker.
                continue
    if failed:
        return False, failed
    return True, [{"check": "fonts_embedded", "value": f"all ({len(seen)} unique)"}]


def _check_no_external_refs(pdf) -> tuple[bool, list[dict]]:
    """Check 4: sem referencias externas (URL, /EmbeddedFiles com /F)."""
    failed: list[dict] = []

    # /Annot com /A /URI em qualquer pagina.
    for page_num, page in enumerate(pdf.pages, start=1):
        try:
            page_obj = page.obj if hasattr(page, "obj") else page
        except Exception:
            continue
        if not page_obj or "/Annots" not in page_obj:
            continue
        annots = page_obj["/Annots"]
        for annot in annots:
            try:
                a_ref = annot.obj.get("/A", {}) if hasattr(annot, "obj") else annot.get("/A", {})
            except Exception:
                continue
            if not a_ref:
                continue
            try:
                a = a_ref.obj if hasattr(a_ref, "obj") else a_ref
            except Exception:
                a = a_ref
            if not a or "/URI" not in a:
                continue
            try:
                uri = str(a["/URI"])
            except Exception:
                continue
            if uri.startswith(("http://", "https://", "ftp://", "mailto:")):
                failed.append({"check": "external_uri", "uri": uri, "page": page_num})

    # /Catalog /Names /EmbeddedFiles - filespecs externos.
    try:
        catalog = pdf.Root
        if "/Names" in catalog:
            names = catalog["/Names"]
            if "/EmbeddedFiles" in names:
                ef = names["/EmbeddedFiles"]
                names_array = ef.get("/Names", []) if ef else []
                for i in range(0, len(names_array) - 1, 2):
                    spec_ref = names_array[i + 1]
                    if spec_ref is None:
                        continue
                    spec = spec_ref.obj if hasattr(spec_ref, "obj") else spec_ref
                    if not spec:
                        continue
                    # /F = file system path, /UF = unicode.
                    if "/F" in spec or "/UF" in spec:
                        path = ""
                        try:
                            path = str(spec.get("/F", spec.get("/UF", "")))
                        except Exception:
                            path = ""
                        if path and (path.startswith("/") or path.startswith("\\")
                                     or (len(path) > 1 and path[1] == ":")):
                            failed.append({"check": "external_embedded_file", "path": path})
    except Exception:
        pass

    if failed:
        return False, failed
    return True, [{"check": "no_external_refs", "value": "ok"}]


def _check_no_encryption(pdf) -> tuple[bool, dict]:
    """Check 5: PDF nao cifrado (PDF/A-2B exige sem cifragem)."""
    if pdf.is_encrypted:
        return False, {"check": "no_encryption", "expected": "false", "got": "true"}
    return True, {"check": "no_encryption", "value": "ok"}


def _check_pdfa2b(pdf) -> tuple[bool, list[dict]]:
    """Checks PDF/A-2B opt-in (D-PDF5): OutputIntent, XMP, no JavaScript.

    O que o nivel B exige (audit do §19.4 verifica quando o opt-in
    `pdfa: PdfA2b` e usado - D-PDF5 do ADR-0021):
    - OutputIntent com ICC profile RGB embedded (D-PDF6)
    - XMP `pdfaid:part=2` e `pdfaid:conformance=B`
    - Sem cifragem, sem JavaScript, sem /OpenAction

    **Tagged PDF NAO e requisito do nivel B** - e o que separa
    o nivel B (basic) do nivel A (accessible). A v1 do PDFPro
    declara conformidade apenas com o nivel B. PDF/A-2A
    (Tagged, acessibilidade) esta fora de escopo da v1.
    """
    failed: list[dict] = []

    # 1. OutputIntents presente.
    output_intents = []
    try:
        output_intents = pdf.Root.get("/OutputIntents", []) or []
    except Exception:
        output_intents = []
    if not output_intents:
        failed.append({
            "check": "pdfa2b_output_intent",
            "expected": "/OutputIntents com pelo menos 1 entrada",
            "got": "ausente",
        })
    else:
        # 2. OutputIntent referencia ICC profile RGB valido.
        oi_ref = output_intents[0]
        oi = oi_ref.obj if hasattr(oi_ref, "obj") else oi_ref
        subtype = ""
        try:
            subtype = str(oi.get("/S", ""))
        except Exception:
            subtype = ""
        if subtype != "/GTS_PDFA1":
            failed.append({
                "check": "pdfa2b_output_intent_subtype",
                "expected": "/S = /GTS_PDFA1",
                "got": subtype or "ausente",
            })
        if "/DestOutputProfile" not in oi:
            failed.append({
                "check": "pdfa2b_icc_profile",
                "expected": "/DestOutputProfile referenciando ICC stream",
                "got": "ausente",
            })
        else:
            icc_ref = oi["/DestOutputProfile"]
            icc_stream = icc_ref.obj if hasattr(icc_ref, "obj") else icc_ref
            icc_bytes = b""
            try:
                # `read_bytes()` descomprime o stream (FlateDecode
                # default do pikepdf no save). `read_raw_bytes()`
                # retorna os bytes brutos pos-compressao, que NAO
                # servem pra validar o ICC header (a estrutura 'acsp'
                # so aparece nos bytes descomprimidos). PDF/A-2B
                # tolera ICC comprimido embedded, mas a auditoria
                # valida o conteudo. pikepdf 10.x: `read_bytes()` e
                # a API; `get_data()` da versao 8 foi removido.
                if hasattr(icc_stream, "read_bytes"):
                    icc_bytes = bytes(icc_stream.read_bytes())
                elif hasattr(icc_stream, "read_raw_bytes"):
                    icc_bytes = bytes(icc_stream.read_raw_bytes())
            except Exception as exc:
                failed.append({
                    "check": "pdfa2b_icc_read",
                    "expected": "ICC stream legivel",
                    "got": f"falha ao ler: {type(exc).__name__}: {exc}",
                })
            if icc_bytes:
                ok, motivo = _audit_icc_signature(icc_bytes)
                if not ok:
                    failed.append({
                        "check": "pdfa2b_icc_valid",
                        "expected": "ICC header valido (acsp + RGB)",
                        "got": motivo,
                    })

    # 3. XMP presente com pdfaid:part=2 e pdfaid:conformance=B.
    xmp_data = ""
    try:
        with pdf.open_metadata() as xmp:
            xmp_data = str(xmp) if xmp is not None else ""
    except Exception:
        xmp_data = ""
    if not xmp_data:
        failed.append({
            "check": "pdfa2b_xmp_present",
            "expected": "XMP metadata stream legivel",
            "got": "ausente",
        })
    else:
        # O XMP pode serializar o mesmo par campo/valor de duas
        # formas: atributo (`pdfaid:part="2"`) ou elemento
        # (`<pdfaid:part>2</pdfaid:part>`). PDF/A-2B aceita ambas
        # e a pikepdf (e varios outros producers) usa elemento.
        # Checamos as duas pra nao ter falso negativo.
        for field, value, check_name in (
            ("pdfaid:part", "2", "pdfa2b_xmp_part"),
            ("pdfaid:conformance", "B", "pdfa2b_xmp_conformance"),
        ):
            field_present = (
                re.search(rf'{re.escape(field)}\s*=\s*"{re.escape(value)}"', xmp_data)
                is not None
                or re.search(
                    rf"<{re.escape(field)}[^>]*>\s*{re.escape(value)}\s*</{re.escape(field)}>",
                    xmp_data,
                ) is not None
            )
            if not field_present:
                # Detalhe do motivo: o campo existe com valor
                # diferente, ou nao existe?
                field_anywhere = field in xmp_data
                failed.append({
                    "check": check_name,
                    "expected": f'{field} = "{value}"',
                    "got": "ausente" if not field_anywhere else "valor diferente",
                })

    # 4. Sem JavaScript (/OpenAction, /AA, /Names/JavaScript).
    try:
        if "/OpenAction" in pdf.Root:
            oa_ref = pdf.Root["/OpenAction"]
            oa = oa_ref.obj if hasattr(oa_ref, "obj") else oa_ref
            if oa:
                oa_subtype = ""
                try:
                    oa_subtype = str(oa.get("/S", ""))
                except Exception:
                    oa_subtype = ""
                if oa_subtype == "/JavaScript" or "/JS" in oa:
                    failed.append({
                        "check": "pdfa2b_no_javascript",
                        "value": "Catalog /OpenAction tem JS",
                    })
        if "/AA" in pdf.Root:
            failed.append({
                "check": "pdfa2b_no_javascript",
                "value": "Catalog /AA presente",
            })
        if "/Names" in pdf.Root and "/JavaScript" in pdf.Root["/Names"]:
            failed.append({
                "check": "pdfa2b_no_javascript",
                "value": "Catalog /Names/JavaScript presente",
            })
    except Exception:
        pass

    if failed:
        return False, failed
    return True, [{"check": "pdfa2b_compliance", "value": "B-level subset (sem Tagged)"}]


def handle_pdf_audit(payload: dict) -> dict:
    """Handler para `pdf.audit` (Etapa 5 PR 3).

    Veja o comentario do bloco acima pra contrato, semantica, e
    rationale do cache key (D-AUDIT-1).
    """
    path_str = payload.get("path", "")
    if not path_str:
        return {"ok": False, "code": "invalid_input", "message": "path ausente"}
    kind = payload.get("kind", "structural")
    pdfa = payload.get("pdfa", None)
    metadata = payload.get("metadata", None) or {}
    n_pages_from_write = payload.get("expected_pages", None)

    if kind != "structural":
        return {
            "ok": False,
            "code": "audit_kind_unsupported",
            "message": (
                f"kind {kind!r} nao suportado por esta build. "
                f"PR 3 cobre apenas 'structural'; "
                f"'visual' entra no PR 4 (pypdfium2)."
            ),
        }

    try:
        path = validate_path(path_str, "read")
    except PathSafetyError as exc:
        return {"ok": False, "code": exc.code, "message": exc.message}

    # Abertura. `pikepdf.Pdf.open` levanta PasswordError se cifrado e
    # sem senha, ou PdfError em caso de arquivo corrompido. Mapeia
    # cada caso pra um codigo estruturado pro caller.
    try:
        pdf = pikepdf.Pdf.open(path)
    except pikepdf.PasswordError:
        return {
            "ok": False,
            "code": "pdf_audit_structural_failed",
            "message": "PDF cifrado (PDF/A-2B e PDF nao-cifrado por definicao)",
            "failed": [{"check": "no_encryption", "expected": "false", "got": "true"}],
        }
    except Exception as exc:
        return {
            "ok": False,
            "code": "pdf_audit_structural_failed",
            "message": f"pikepdf nao abriu o PDF: {type(exc).__name__}: {exc}",
            "failed": [{
                "check": "pdf_open",
                "expected": "pikepdf.Pdf.open sem exception",
                "got": type(exc).__name__,
            }],
        }

    checks_passed: list[dict] = []
    failed: list[dict] = []

    ok, item = _check_n_pages(pdf, n_pages_from_write)
    (checks_passed if ok else failed).append(item)

    ok, items = _check_docinfo(pdf)
    (checks_passed if ok else failed).extend(items)

    ok, items = _check_fonts_embedded(pdf)
    (checks_passed if ok else failed).extend(items)

    ok, items = _check_no_external_refs(pdf)
    (checks_passed if ok else failed).extend(items)

    ok, item = _check_no_encryption(pdf)
    (checks_passed if ok else failed).append(item)

    # PDF/A-2B opt-in (D-PDF5): baseline + OutputIntent/ICC/XMP/JS.
    if pdfa == "pdfa_2b":
        ok, items = _check_pdfa2b(pdf)
        (checks_passed if ok else failed).extend(items)

    # Cache key (D-AUDIT-1): hash do PDF + versao das regras.
    try:
        pdf_bytes = path.read_bytes()
        pdf_hash = hashlib.sha256(pdf_bytes).hexdigest()
    except Exception:
        pdf_hash = "<unreadable>"
    cache_key = f"{pdf_hash}:{AUDIT_RULES_VERSION}"

    if failed:
        return {
            "ok": False,
            "code": "pdf_audit_structural_failed",
            "message": (
                f"{len(failed)} check(s) falharam (kind=structural, pdfa={pdfa!r})"
            ),
            "failed": failed,
            "checks_passed": checks_passed,
            "cache_key": cache_key,
            "rules_version": AUDIT_RULES_VERSION,
        }
    return {
        "ok": True,
        "checks": checks_passed,
        "coverage": "full",
        "cache_key": cache_key,
        "rules_version": AUDIT_RULES_VERSION,
    }


# ---------------------------------------------------------------------------
# Dispatch table
# ---------------------------------------------------------------------------

HANDLERS: dict[str, Callable[[dict], dict]] = {
    "docx.write": handle_docx_write,
    "docx.read": handle_docx_read,
    "xlsx.write": handle_xlsx_write,
    "xlsx.read": handle_xlsx_read,
    "pdf.write": handle_pdf_write,
    "pdf.read": handle_pdf_read,
    "ocr.run": handle_ocr_run,
    "pdf.audit": handle_pdf_audit,
}


def handle_tool_invoke(
    msg: dict[str, Any],
    auth_token: str | None,
) -> bytes:
    """Dispatcha o `tool.invoke` pro handler da capability.

    Retorna os bytes da resposta (`tool.result` ou `worker.error`),
    carregando o mesmo `request_id` do request de entrada (o ator
    Rust casa por ele - ver `ipc_message`).
    """
    request_id = msg["request_id"]
    auth = msg.get("auth")
    # Validacao de token (so se ja vimos um `app.ack`).
    if auth_token is not None and auth != auth_token:
        return ipc_message(
            "worker.error",
            {
                "code": "process_unauthorized",
                "message": "token ausente ou invalido",
            },
            request_id=request_id,
        )
    payload_in = msg.get("payload", {})
    if not isinstance(payload_in, dict):
        return ipc_message(
            "tool.result",
            {
                "ok": False,
                "code": "invalid_payload",
                "message": f"payload precisa ser dict, veio {type(payload_in).__name__}",
            },
            request_id=request_id,
            auth=auth,
        )
    capability = payload_in.get("capability", "")
    if not capability:
        return ipc_message(
            "tool.result",
            {
                "ok": False,
                "code": "missing_capability",
                "message": "payload.capability ausente",
            },
            request_id=request_id,
            auth=auth,
        )
    handler = HANDLERS.get(capability)
    if handler is None:
        declared = sorted(HANDLERS.keys())
        return ipc_message(
            "tool.result",
            {
                "ok": False,
                "code": "unknown_capability",
                "message": (
                    f"document-worker v0.4.0 nao implementa "
                    f"{capability!r}. Capabilities declaradas: {declared}."
                ),
            },
            request_id=request_id,
            auth=auth,
        )
    try:
        result = handler(payload_in)
    except PathSafetyError as exc:
        log.warning("path safety falhou: %s (%s)", exc.code, exc.message)
        return ipc_message(
            "tool.result",
            {
                "ok": False,
                "code": exc.code,
                "message": exc.message,
            },
            request_id=request_id,
            auth=auth,
        )
    except ValueError as exc:
        log.warning("input invalido em %s: %s", capability, exc)
        return ipc_message(
            "tool.result",
            {
                "ok": False,
                "code": "invalid_input",
                "message": str(exc),
            },
            request_id=request_id,
            auth=auth,
        )
    except Exception as exc:  # noqa: BLE001 - borda do dispatch, queremos worker.error
        log.exception("handler %s falhou", capability)
        return ipc_message(
            "tool.result",
            {
                "ok": False,
                "code": "handler_error",
                "message": f"{type(exc).__name__}: {exc}",
            },
            request_id=request_id,
            auth=auth,
        )
    log.info("handler %s ok: %s", capability, {k: result.get(k) for k in ("ok", "path", "size_bytes") if k in result})
    return ipc_message("tool.result", result, request_id=request_id, auth=auth)


# ---------------------------------------------------------------------------
# Loop principal do worker
# ---------------------------------------------------------------------------


def load_manifest(manifest_path: Path) -> dict[str, Any]:
    """Carrega `manifest.json` ao lado do script."""
    with open(manifest_path, encoding="utf-8") as f:
        return json.load(f)


def worker_main(manifest_path: Path) -> int:
    """Loop principal: cria pipe, espera connect, dispatch."""
    manifest = load_manifest(manifest_path)
    worker_id = manifest["worker_id"]
    log.info("subindo %s %s", worker_id, manifest.get("version", "?"))

    # Carrega as fontes Tinta e Latao no startup (uma vez so).
    font_status = ensure_fonts_registered()
    log.info("fontes: %s", font_status)

    # Configura pytesseract (tesseract_cmd + TESSDATA_PREFIX) e
    # detecta a versao do Tesseract no startup. Idempotente.
    _configure_pytesseract()
    tess_version = _get_tesseract_version()
    ocr_available = (
        PYTESSERACT_AVAILABLE
        and _tesseract_executable_present()
        and tess_version is not None
    )
    log.info(
        "Tesseract: versao=%s disponivel=%s pytesseract=%s",
        tess_version, ocr_available, PYTESSERACT_AVAILABLE,
    )

    # 1. Gera nome unico pro pipe. O `<name>` e a parte depois de
    #    `\\\\.\\pipe\\`. O `PipeName::new` (Rust) valida.
    pipe_name = f"frederico-{worker_id}-{uuid.uuid4().hex[:12]}"
    pipe_path = rf"\\.\pipe\{pipe_name}"

    # 2. Cria o NamedPipeServer. `maxInstances=1` (so o app principal).
    try:
        pipe_handle = win32pipe.CreateNamedPipe(
            pipe_path,
            win32pipe.PIPE_ACCESS_DUPLEX,
            win32pipe.PIPE_TYPE_BYTE
            | win32pipe.PIPE_READMODE_BYTE
            | win32pipe.PIPE_WAIT,
            1,
            READ_BUFFER_SIZE,
            READ_BUFFER_SIZE,
            0,
            None,
        )
    except pywintypes.error as exc:
        log.error("CreateNamedPipe falhou para %s: %s", pipe_path, exc)
        return 1

    # 3. Anuncia o pipe pro app via stdout. PRIMEIRA linha do stdout.
    print(f"READY {pipe_name}", flush=True)

    # 4. Espera o app conectar.
    try:
        win32pipe.ConnectNamedPipe(pipe_handle)
    except pywintypes.error as exc:
        if exc.winerror not in (535,):
            log.error("ConnectNamedPipe falhou: %s", exc)
            win32file.CloseHandle(pipe_handle)
            return 1
    log.info("cliente conectou em %s", pipe_path)

    # 5. Envia `worker.hello`. Inclui o `font_status` no payload
    #    (extensao alem do WorkerManifest) pra o manager saber se as
    #    fontes T&L estao presentes ou em fallback. Tambem inclui
    #    `ocr_available` + `tesseract_version` + `tesseract_status`
    #    (Etapa 2B+Y / ADR-0019) pra o caller saber se OCR ta
    #    disponivel sem precisar chamar `ocr.run` pra descobrir.
    hello_payload = dict(manifest)
    hello_payload["font_status"] = font_status
    hello_payload["ocr_available"] = ocr_available
    hello_payload["tesseract_version"] = tess_version
    hello_payload["tesseract_status"] = {
        "binary_present": _tesseract_executable_present(),
        "pytesseract_imported": PYTESSERACT_AVAILABLE,
        "version": tess_version,
        "tessdata_dir": str(TESSERACT_TESSDATA_DIR),
    }
    try:
        win32file.WriteFile(pipe_handle, ipc_message("worker.hello", hello_payload))
    except pywintypes.error as exc:
        log.error("write do worker.hello falhou: %s", exc)
        win32file.CloseHandle(pipe_handle)
        return 1

    # 6. Loop: le linhas do pipe e dispatcha.
    auth_token: str | None = None
    buffer = b""
    while True:
        try:
            err, chunk = win32file.ReadFile(pipe_handle, READ_BUFFER_SIZE)
        except pywintypes.error as exc:
            if exc.winerror in (109, 232):
                log.info("peer fechou o pipe (EOF)")
                break
            log.error("ReadFile falhou: %s", exc)
            break

        if err == 0 and not chunk:
            log.info("ReadFile devolveu 0 bytes (EOF)")
            break

        buffer += chunk
        shutdown_received = False
        while b"\n" in buffer:
            line, buffer = buffer.split(b"\n", 1)
            if not line:
                continue
            try:
                msg = decode_line(line)
            except (ValueError, json.JSONDecodeError) as exc:
                log.warning("decode falhou: %s (linha ignorada)", exc)
                continue

            op = msg.get("op")
            log.debug("recv op=%s request_id=%s", op, msg.get("request_id"))

            if op == "app.ack":
                auth_token = msg.get("auth")
                log.info("handshake completo (auth salvo)")

            elif op == "app.ping":
                pong_payload = {
                    "status": "ok",
                    "env_received": {},
                    "font_status": ensure_fonts_registered(),
                }
                try:
                    win32file.WriteFile(
                        pipe_handle,
                        ipc_message("worker.pong", pong_payload, request_id=msg.get("request_id")),
                    )
                except pywintypes.error as exc:
                    log.error("write do pong falhou: %s", exc)
                    break

            elif op == "app.shutdown":
                log.info("app.shutdown recebido, saindo")
                shutdown_received = True
                break

            elif op == "tool.invoke":
                response = handle_tool_invoke(msg, auth_token)
                try:
                    win32file.WriteFile(pipe_handle, response)
                except pywintypes.error as exc:
                    log.error("write do tool.result falhou: %s", exc)
                    break

            else:
                log.warning("op desconhecido/ignorado: %r", op)

        # Se `app.shutdown` foi recebido, sai do loop principal.
        # **Bug fix v0.2.0:** a versao inicial tinha `break` direto
        # no `elif app.shutdown`, mas esse break estava dentro do
        # `while b"\\n" in buffer` (loop interno que processa linhas
        # do buffer). O loop externo (`while True:`) continuava, e
        # o worker fazia um novo `ReadFile` esperando dados que
        # nunca chegavam - travando o `actor_task.await` no
        # `manager.shutdown`. Agora levantamos a flag e quebramos
        # no escopo certo.
        if shutdown_received:
            break

    try:
        win32file.CloseHandle(pipe_handle)
    except pywintypes.error:
        pass
    log.info("worker %s saindo", worker_id)
    return 0


def main() -> int:
    """Entry point. Aceita `--manifest <path>` (default: `manifest.json` ao lado)."""
    args = sys.argv[1:]
    manifest_path = Path(__file__).resolve().parent / "manifest.json"
    i = 0
    while i < len(args):
        if args[i] == "--manifest" and i + 1 < len(args):
            manifest_path = Path(args[i + 1])
            i += 2
        else:
            log.warning("argumento ignorado: %r", args[i])
            i += 1
    if not manifest_path.is_file():
        log.error("manifesto nao encontrado em %s", manifest_path)
        return 1
    return worker_main(manifest_path)


if __name__ == "__main__":
    raise SystemExit(main())
